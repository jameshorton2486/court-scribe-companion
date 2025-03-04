
import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext
import os
import re
import docx
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import time
import threading
import openai
import io
from PIL import Image, ImageTk
import matplotlib.pyplot as plt
import base64
import requests
from matplotlib.backends.backend_agg import FigureCanvasAgg

class BookProcessor:
    def __init__(self, root):
        self.root = root
        self.root.title("Book Processor")
        self.root.geometry("1000x800")
        self.root.configure(bg="#f0f0f0")
        
        # Setup variables
        self.input_file = tk.StringVar()
        self.output_dir = tk.StringVar()
        self.book_title = tk.StringVar()
        self.author_name = tk.StringVar()
        self.current_status = tk.StringVar(value="Ready")
        self.progress_value = tk.DoubleVar(value=0.0)
        self.openai_api_key = tk.StringVar()
        self.chapter_outline = tk.StringVar()
        
        # Default output directory is 'output' in current directory
        default_output = os.path.join(os.getcwd(), "output")
        self.output_dir.set(default_output)
        
        # Document storage
        self.docx_content = None
        self.chapters = []
        self.toc = []
        self.current_chapter_index = 0
        
        # Create the main frame
        self.create_widgets()
    
    def create_widgets(self):
        # Main frame with notebook (tabs)
        main_frame = ttk.Frame(self.root, padding=10)
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        notebook = ttk.Notebook(main_frame)
        notebook.pack(fill=tk.BOTH, expand=True)
        
        # Tab 1: File Processing
        file_tab = ttk.Frame(notebook, padding=10)
        notebook.add(file_tab, text="File Processing")
        
        # Tab 2: AI Enhancement
        ai_tab = ttk.Frame(notebook, padding=10)
        notebook.add(ai_tab, text="AI Enhancement")
        
        # Tab 3: Chapter Generation
        chapter_tab = ttk.Frame(notebook, padding=10)
        notebook.add(chapter_tab, text="Chapter Generation")
        
        # File tab contents
        self.create_file_tab(file_tab)
        
        # AI Enhancement tab contents
        self.create_ai_tab(ai_tab)
        
        # Chapter Generation tab contents
        self.create_chapter_tab(chapter_tab)
        
        # Log display (shared across tabs)
        log_frame = ttk.LabelFrame(main_frame, text="Processing Log", padding=10)
        log_frame.pack(fill=tk.BOTH, expand=True, pady=5)
        
        self.log_text = scrolledtext.ScrolledText(log_frame, width=80, height=10)
        self.log_text.pack(fill=tk.BOTH, expand=True)
        self.log_text.config(state=tk.DISABLED)
        
        # Progress and status (shared across tabs)
        progress_frame = ttk.Frame(main_frame)
        progress_frame.pack(fill=tk.X, pady=5)
        
        ttk.Label(progress_frame, text="Status:").pack(side=tk.LEFT, padx=5)
        ttk.Label(progress_frame, textvariable=self.current_status).pack(side=tk.LEFT, padx=5)
        
        self.progress_bar = ttk.Progressbar(progress_frame, variable=self.progress_value, length=400, mode="determinate")
        self.progress_bar.pack(side=tk.RIGHT, padx=5)
    
    def create_file_tab(self, parent):
        # File input section
        file_frame = ttk.LabelFrame(parent, text="File Selection", padding=10)
        file_frame.pack(fill=tk.X, pady=5)
        
        ttk.Label(file_frame, text="Input DOCX File:").grid(row=0, column=0, sticky=tk.W, pady=5)
        ttk.Entry(file_frame, textvariable=self.input_file, width=50).grid(row=0, column=1, padx=5, pady=5)
        ttk.Button(file_frame, text="Browse...", command=self.browse_input_file).grid(row=0, column=2, padx=5, pady=5)
        
        ttk.Label(file_frame, text="Output Directory:").grid(row=1, column=0, sticky=tk.W, pady=5)
        ttk.Entry(file_frame, textvariable=self.output_dir, width=50).grid(row=1, column=1, padx=5, pady=5)
        ttk.Button(file_frame, text="Browse...", command=self.browse_output_dir).grid(row=1, column=2, padx=5, pady=5)
        
        # Book metadata section
        metadata_frame = ttk.LabelFrame(parent, text="Book Metadata", padding=10)
        metadata_frame.pack(fill=tk.X, pady=5)
        
        ttk.Label(metadata_frame, text="Book Title:").grid(row=0, column=0, sticky=tk.W, pady=5)
        ttk.Entry(metadata_frame, textvariable=self.book_title, width=50).grid(row=0, column=1, padx=5, pady=5)
        
        ttk.Label(metadata_frame, text="Author Name:").grid(row=1, column=0, sticky=tk.W, pady=5)
        ttk.Entry(metadata_frame, textvariable=self.author_name, width=50).grid(row=1, column=1, padx=5, pady=5)
        
        # Processing options section
        options_frame = ttk.LabelFrame(parent, text="Processing Options", padding=10)
        options_frame.pack(fill=tk.X, pady=5)
        
        # Add checkboxes for different processing options
        self.fix_encoding = tk.BooleanVar(value=True)
        ttk.Checkbutton(options_frame, text="Fix Text Encoding Issues", variable=self.fix_encoding).grid(row=0, column=0, sticky=tk.W, pady=2)
        
        self.generate_toc = tk.BooleanVar(value=True)
        ttk.Checkbutton(options_frame, text="Generate Table of Contents", variable=self.generate_toc).grid(row=1, column=0, sticky=tk.W, pady=2)
        
        self.create_chapters = tk.BooleanVar(value=True)
        ttk.Checkbutton(options_frame, text="Create Chapter Outlines", variable=self.create_chapters).grid(row=2, column=0, sticky=tk.W, pady=2)
        
        self.enhance_content = tk.BooleanVar(value=False)
        ttk.Checkbutton(options_frame, text="Enhance Book Content (Experimental)", variable=self.enhance_content).grid(row=3, column=0, sticky=tk.W, pady=2)
        
        # Action buttons
        button_frame = ttk.Frame(parent)
        button_frame.pack(fill=tk.X, pady=10)
        
        ttk.Button(button_frame, text="Load Document", command=self.load_document, width=20).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="Process Document", command=self.process_document, width=20).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="Save All Chapters", command=self.save_all_chapters, width=20).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="Generate Complete Book", command=self.generate_complete_book, width=20).pack(side=tk.LEFT, padx=5)
    
    def create_ai_tab(self, parent):
        # OpenAI API Key
        api_frame = ttk.LabelFrame(parent, text="OpenAI Configuration", padding=10)
        api_frame.pack(fill=tk.X, pady=5)
        
        ttk.Label(api_frame, text="OpenAI API Key:").grid(row=0, column=0, sticky=tk.W, pady=5)
        api_entry = ttk.Entry(api_frame, textvariable=self.openai_api_key, width=50, show="*")
        api_entry.grid(row=0, column=1, padx=5, pady=5)
        
        # Toggle button to show/hide API key
        self.show_api_key = tk.BooleanVar(value=False)
        def toggle_api_key_visibility():
            if self.show_api_key.get():
                api_entry.config(show="")
            else:
                api_entry.config(show="*")
        
        ttk.Checkbutton(api_frame, text="Show API Key", variable=self.show_api_key, command=toggle_api_key_visibility).grid(row=0, column=2, padx=5, pady=5)
        
        # Test API key button
        ttk.Button(api_frame, text="Test API Connection", command=self.test_openai_connection).grid(row=1, column=1, padx=5, pady=5)
        
        # AI Review options
        review_frame = ttk.LabelFrame(parent, text="AI Content Review", padding=10)
        review_frame.pack(fill=tk.X, pady=5)
        
        # AI Review options
        self.review_grammar = tk.BooleanVar(value=True)
        ttk.Checkbutton(review_frame, text="Check Grammar and Style", variable=self.review_grammar).grid(row=0, column=0, sticky=tk.W, pady=2)
        
        self.review_coherence = tk.BooleanVar(value=True)
        ttk.Checkbutton(review_frame, text="Check Content Coherence", variable=self.review_coherence).grid(row=1, column=0, sticky=tk.W, pady=2)
        
        self.suggest_improvements = tk.BooleanVar(value=True)
        ttk.Checkbutton(review_frame, text="Suggest Content Improvements", variable=self.suggest_improvements).grid(row=2, column=0, sticky=tk.W, pady=2)
        
        self.generate_toc_ai = tk.BooleanVar(value=True)
        ttk.Checkbutton(review_frame, text="Generate AI-Assisted Table of Contents", variable=self.generate_toc_ai).grid(row=3, column=0, sticky=tk.W, pady=2)
        
        # Action buttons
        ai_button_frame = ttk.Frame(parent)
        ai_button_frame.pack(fill=tk.X, pady=10)
        
        ttk.Button(ai_button_frame, text="Review Content with AI", command=self.review_with_ai, width=25).pack(side=tk.LEFT, padx=5)
        ttk.Button(ai_button_frame, text="Generate AI Table of Contents", command=self.generate_ai_toc, width=25).pack(side=tk.LEFT, padx=5)
    
    def create_chapter_tab(self, parent):
        # Chapter selection section
        chapter_select_frame = ttk.LabelFrame(parent, text="Chapter Selection", padding=10)
        chapter_select_frame.pack(fill=tk.X, pady=5)
        
        self.chapter_list = tk.StringVar()
        self.chapter_listbox = tk.Listbox(chapter_select_frame, height=5, width=80)
        self.chapter_listbox.pack(fill=tk.X, padx=5, pady=5)
        self.chapter_listbox.bind('<<ListboxSelect>>', self.on_chapter_select)
        
        # Navigation buttons
        nav_frame = ttk.Frame(chapter_select_frame)
        nav_frame.pack(fill=tk.X, pady=5)
        
        ttk.Button(nav_frame, text="Previous Chapter", command=self.prev_chapter).pack(side=tk.LEFT, padx=5)
        ttk.Button(nav_frame, text="Next Chapter", command=self.next_chapter).pack(side=tk.LEFT, padx=5)
        
        # Chapter outline section
        outline_frame = ttk.LabelFrame(parent, text="Chapter Outline", padding=10)
        outline_frame.pack(fill=tk.BOTH, expand=True, pady=5)
        
        self.outline_text = scrolledtext.ScrolledText(outline_frame, width=80, height=5)
        self.outline_text.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        # Chapter generation options
        gen_options_frame = ttk.LabelFrame(parent, text="Generation Options", padding=10)
        gen_options_frame.pack(fill=tk.X, pady=5)
        
        self.include_images = tk.BooleanVar(value=True)
        ttk.Checkbutton(gen_options_frame, text="Include Images in Generated Content", variable=self.include_images).grid(row=0, column=0, sticky=tk.W, pady=2)
        
        self.include_diagrams = tk.BooleanVar(value=True)
        ttk.Checkbutton(gen_options_frame, text="Include Diagrams/Charts", variable=self.include_diagrams).grid(row=1, column=0, sticky=tk.W, pady=2)
        
        self.writing_style = tk.StringVar(value="informative")
        ttk.Label(gen_options_frame, text="Writing Style:").grid(row=2, column=0, sticky=tk.W, pady=5)
        style_combo = ttk.Combobox(gen_options_frame, textvariable=self.writing_style, state="readonly")
        style_combo['values'] = ('informative', 'academic', 'conversational', 'narrative', 'technical')
        style_combo.grid(row=2, column=1, padx=5, pady=5)
        
        # Generation buttons
        gen_button_frame = ttk.Frame(parent)
        gen_button_frame.pack(fill=tk.X, pady=10)
        
        ttk.Button(gen_button_frame, text="Generate Chapter Content", command=self.generate_chapter_content, width=25).pack(side=tk.LEFT, padx=5)
        ttk.Button(gen_button_frame, text="Save Generated Content", command=self.save_generated_chapter, width=25).pack(side=tk.LEFT, padx=5)
        
        # Preview frame for generated content
        preview_frame = ttk.LabelFrame(parent, text="Content Preview", padding=10)
        preview_frame.pack(fill=tk.BOTH, expand=True, pady=5)
        
        self.preview_text = scrolledtext.ScrolledText(preview_frame, width=80, height=10)
        self.preview_text.pack(fill=tk.BOTH, expand=True)
        
        # Image preview
        self.image_preview_frame = ttk.Frame(preview_frame)
        self.image_preview_frame.pack(fill=tk.X, pady=5)
        self.image_label = ttk.Label(self.image_preview_frame)
        self.image_label.pack(pady=5)
    
    def browse_input_file(self):
        file_path = filedialog.askopenfilename(
            filetypes=[("Word Documents", "*.docx"), ("All Files", "*.*")]
        )
        if file_path:
            self.input_file.set(file_path)
            # Try to extract title from filename
            filename = os.path.basename(file_path)
            title, _ = os.path.splitext(filename)
            if not self.book_title.get():
                self.book_title.set(title)
    
    def browse_output_dir(self):
        dir_path = filedialog.askdirectory()
        if dir_path:
            self.output_dir.set(dir_path)
    
    def log(self, message):
        self.log_text.config(state=tk.NORMAL)
        self.log_text.insert(tk.END, f"{message}\n")
        self.log_text.see(tk.END)
        self.log_text.config(state=tk.DISABLED)
        self.root.update()
    
    def update_progress(self, value, status=None):
        self.progress_value.set(value)
        if status:
            self.current_status.set(status)
        self.root.update()
    
    def load_document(self):
        if not self.input_file.get():
            messagebox.showerror("Error", "Please select an input file")
            return
        
        try:
            self.log(f"Loading document: {self.input_file.get()}")
            self.update_progress(10, "Loading document...")
            
            # Load the document
            doc = Document(self.input_file.get())
            self.docx_content = doc
            
            self.update_progress(50, "Extracting text...")
            # Extract text
            full_text = ""
            for para in doc.paragraphs:
                full_text += para.text + "\n"
            
            self.log(f"Document loaded successfully. {len(doc.paragraphs)} paragraphs found.")
            
            # Set book title if not already set
            if not self.book_title.get():
                # Try to find title in the first few paragraphs
                for para in doc.paragraphs[:10]:
                    if para.style.name.startswith('Heading') or para.style.name.startswith('Title'):
                        self.book_title.set(para.text)
                        break
            
            self.update_progress(100, "Document loaded successfully")
            messagebox.showinfo("Success", "Document loaded successfully")
            
        except Exception as e:
            self.log(f"Error loading document: {str(e)}")
            messagebox.showerror("Error", f"Failed to load document: {str(e)}")
            self.update_progress(0, "Error loading document")
    
    def process_document(self):
        if not self.docx_content:
            messagebox.showerror("Error", "Please load a document first")
            return
        
        # Start processing in a separate thread to avoid freezing the UI
        threading.Thread(target=self._process_document_thread, daemon=True).start()
    
    def _process_document_thread(self):
        # ... keep existing code (document processing functionality)
    
    def fix_text_encoding(self):
        # ... keep existing code (text encoding fix functionality)
    
    def enhance_book_content(self):
        # ... keep existing code (book enhancement functionality)
    
    def save_all_chapters(self):
        # ... keep existing code (chapter saving functionality)
    
    def generate_complete_book(self):
        # ... keep existing code (complete book generation functionality)
    
    def test_openai_connection(self):
        api_key = self.openai_api_key.get()
        if not api_key:
            messagebox.showerror("Error", "Please enter an OpenAI API key")
            return
        
        self.update_progress(0, "Testing OpenAI connection...")
        
        try:
            # Set the API key
            openai.api_key = api_key
            
            # Simple test request
            response = openai.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=[
                    {"role": "system", "content": "You are a helpful assistant."},
                    {"role": "user", "content": "Hello, this is a test message."}
                ],
                max_tokens=20
            )
            
            self.log("OpenAI connection test successful!")
            self.update_progress(100, "OpenAI connection successful")
            messagebox.showinfo("Success", "Successfully connected to OpenAI API")
            
        except Exception as e:
            self.log(f"OpenAI connection error: {str(e)}")
            self.update_progress(0, "OpenAI connection failed")
            messagebox.showerror("Error", f"Failed to connect to OpenAI: {str(e)}")
    
    def review_with_ai(self):
        if not self.chapters:
            messagebox.showerror("Error", "No chapters found. Please process a document first.")
            return
        
        api_key = self.openai_api_key.get()
        if not api_key:
            messagebox.showerror("Error", "Please enter an OpenAI API key")
            return
        
        openai.api_key = api_key
        
        # Start AI review in a separate thread
        threading.Thread(target=self._review_with_ai_thread, daemon=True).start()
    
    def _review_with_ai_thread(self):
        try:
            self.log("Starting AI content review...")
            self.update_progress(0, "Starting AI review...")
            
            total_chapters = len(self.chapters)
            review_results = []
            
            for i, chapter in enumerate(self.chapters):
                self.update_progress((i / total_chapters) * 100, f"Reviewing chapter {i+1} of {total_chapters}")
                self.log(f"Analyzing chapter: {chapter['title']}")
                
                # Extract all text from this chapter
                chapter_text = ""
                for para in chapter["content"]:
                    chapter_text += para.text + "\n"
                
                # Build prompt based on selected options
                prompt = f"Review this book chapter titled '{chapter['title']}'. "
                
                if self.review_grammar.get():
                    prompt += "Check for grammar and style issues. "
                
                if self.review_coherence.get():
                    prompt += "Evaluate content coherence and logical flow. "
                
                if self.suggest_improvements.get():
                    prompt += "Suggest specific improvements to enhance the content. "
                
                # Send to OpenAI
                response = openai.chat.completions.create(
                    model="gpt-4",
                    messages=[
                        {"role": "system", "content": "You are a professional book editor with expertise in improving book content."},
                        {"role": "user", "content": prompt + "\n\nCHAPTER CONTENT:\n" + chapter_text}
                    ],
                    max_tokens=1000
                )
                
                # Store results
                ai_feedback = response.choices[0].message.content
                review_results.append({
                    "chapter_title": chapter["title"],
                    "feedback": ai_feedback
                })
                
                self.log(f"AI review completed for chapter: {chapter['title']}")
            
            # Create review document
            self.update_progress(90, "Creating review document...")
            
            review_doc = Document()
            review_doc.add_heading("AI Content Review", level=1)
            review_doc.add_paragraph(f"Book: {self.book_title.get()}")
            review_doc.add_paragraph(f"Author: {self.author_name.get()}")
            review_doc.add_paragraph(f"Generated on: {time.strftime('%Y-%m-%d %H:%M:%S')}")
            
            for result in review_results:
                review_doc.add_heading(f"Chapter: {result['chapter_title']}", level=2)
                review_doc.add_paragraph(result["feedback"])
                review_doc.add_paragraph()  # Add space between reviews
            
            # Save review document
            os.makedirs(self.output_dir.get(), exist_ok=True)
            review_filepath = os.path.join(self.output_dir.get(), "AI_Review.docx")
            review_doc.save(review_filepath)
            
            self.update_progress(100, "AI review completed")
            self.log(f"AI review completed. Review document saved to {review_filepath}")
            messagebox.showinfo("Success", f"AI review completed. Review document saved to {review_filepath}")
            
        except Exception as e:
            self.log(f"Error during AI review: {str(e)}")
            self.update_progress(0, "Error during AI review")
            messagebox.showerror("Error", f"Failed to complete AI review: {str(e)}")
    
    def generate_ai_toc(self):
        if not self.chapters:
            messagebox.showerror("Error", "No chapters found. Please process a document first.")
            return
        
        api_key = self.openai_api_key.get()
        if not api_key:
            messagebox.showerror("Error", "Please enter an OpenAI API key")
            return
        
        openai.api_key = api_key
        
        # Start TOC generation in a separate thread
        threading.Thread(target=self._generate_ai_toc_thread, daemon=True).start()
    
    def _generate_ai_toc_thread(self):
        try:
            self.log("Generating AI-assisted table of contents...")
            self.update_progress(0, "Starting TOC generation...")
            
            # Extract chapter titles and brief content
            chapters_info = []
            for chapter in self.chapters:
                # Get first 3 paragraphs of content (or fewer if there aren't 3)
                preview = ""
                for i, para in enumerate(chapter["content"][:3]):
                    preview += para.text + "\n"
                
                chapters_info.append({
                    "title": chapter["title"],
                    "preview": preview
                })
            
            # Create prompt for OpenAI
            chapters_text = ""
            for i, ch in enumerate(chapters_info):
                chapters_text += f"Chapter {i+1}: {ch['title']}\nPreview: {ch['preview']}\n\n"
            
            prompt = f"""Based on the following chapter titles and content previews, generate a comprehensive table of contents for the book titled "{self.book_title.get()}".
            For each chapter, add 3-5 detailed subsections that would logically organize the content.
            Format the output as follows:
            
            Chapter 1: [Chapter Title]
            1.1 [Subsection title]
            1.2 [Subsection title]
            ...and so on
            
            Chapters:
            {chapters_text}
            """
            
            # Send to OpenAI
            self.update_progress(30, "Querying OpenAI...")
            response = openai.chat.completions.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": "You are a professional book editor with expertise in organizing book content."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=2000
            )
            
            # Get response
            toc_content = response.choices[0].message.content
            
            # Create TOC document
            self.update_progress(70, "Creating TOC document...")
            
            toc_doc = Document()
            toc_doc.add_heading(f"Table of Contents: {self.book_title.get()}", level=1)
            toc_doc.add_paragraph(f"Author: {self.author_name.get()}")
            toc_doc.add_paragraph(f"Generated on: {time.strftime('%Y-%m-%d %H:%M:%S')}")
            toc_doc.add_paragraph()
            
            # Add the TOC content
            for line in toc_content.split('\n'):
                if line.strip():
                    # Format based on if it's a chapter or subsection
                    if line.startswith('Chapter'):
                        p = toc_doc.add_paragraph()
                        p.add_run(line).bold = True
                    else:
                        toc_doc.add_paragraph(line, style='List Bullet')
            
            # Save TOC document
            os.makedirs(self.output_dir.get(), exist_ok=True)
            toc_filepath = os.path.join(self.output_dir.get(), "AI_Table_of_Contents.docx")
            toc_doc.save(toc_filepath)
            
            # Also update application's TOC
            self.toc = []
            current_chapter = None
            
            for line in toc_content.split('\n'):
                if line.startswith('Chapter'):
                    # This is a main chapter
                    chapter_parts = line.split(':', 1)
                    if len(chapter_parts) > 1:
                        chapter_number = chapter_parts[0].replace('Chapter', '').strip()
                        chapter_title = chapter_parts[1].strip()
                        
                        try:
                            idx = int(chapter_number) - 1
                            if 0 <= idx < len(self.chapters):
                                current_chapter = idx
                                self.toc.append({
                                    "title": chapter_title,
                                    "level": 1,
                                    "index": idx
                                })
                        except ValueError:
                            pass
                        
                elif current_chapter is not None and line.strip():
                    # This is a subsection
                    self.toc.append({
                        "title": line.strip(),
                        "level": 2,
                        "index": current_chapter
                    })
            
            # Populate chapter listbox
            self.chapter_listbox.delete(0, tk.END)
            for i, chapter in enumerate(self.chapters):
                self.chapter_listbox.insert(tk.END, f"Chapter {i+1}: {chapter['title']}")
            
            self.update_progress(100, "TOC generation completed")
            self.log(f"AI TOC generation completed. TOC document saved to {toc_filepath}")
            messagebox.showinfo("Success", f"AI TOC generation completed. Document saved to {toc_filepath}")
            
        except Exception as e:
            self.log(f"Error during TOC generation: {str(e)}")
            self.update_progress(0, "Error during TOC generation")
            messagebox.showerror("Error", f"Failed to generate TOC: {str(e)}")
    
    def on_chapter_select(self, event):
        if not self.chapters:
            return
            
        selection = self.chapter_listbox.curselection()
        if selection:
            index = selection[0]
            self.current_chapter_index = index
            
            # Display chapter title and update outline textarea
            chapter = self.chapters[index]
            self.outline_text.delete(1.0, tk.END)
            self.outline_text.insert(tk.END, f"Outline for: {chapter['title']}\n\n")
            
            # Clear preview
            self.preview_text.delete(1.0, tk.END)
            
            # Remove any image preview
            self.image_label.config(image=None)
    
    def prev_chapter(self):
        if self.chapters and self.current_chapter_index > 0:
            self.current_chapter_index -= 1
            self.chapter_listbox.selection_clear(0, tk.END)
            self.chapter_listbox.selection_set(self.current_chapter_index)
            self.chapter_listbox.see(self.current_chapter_index)
            self.on_chapter_select(None)
    
    def next_chapter(self):
        if self.chapters and self.current_chapter_index < len(self.chapters) - 1:
            self.current_chapter_index += 1
            self.chapter_listbox.selection_clear(0, tk.END)
            self.chapter_listbox.selection_set(self.current_chapter_index)
            self.chapter_listbox.see(self.current_chapter_index)
            self.on_chapter_select(None)
    
    def generate_chapter_content(self):
        if not self.chapters:
            messagebox.showerror("Error", "No chapters available. Please process a document first.")
            return
        
        if self.current_chapter_index < 0 or self.current_chapter_index >= len(self.chapters):
            messagebox.showerror("Error", "Please select a valid chapter first.")
            return
        
        outline = self.outline_text.get(1.0, tk.END).strip()
        if not outline:
            messagebox.showerror("Error", "Please provide an outline for the chapter.")
            return
        
        api_key = self.openai_api_key.get()
        if not api_key:
            messagebox.showerror("Error", "Please enter an OpenAI API key")
            return
        
        openai.api_key = api_key
        
        # Start chapter generation in a separate thread
        threading.Thread(target=self._generate_chapter_thread, daemon=True).start()
    
    def _generate_chapter_thread(self):
        try:
            chapter = self.chapters[self.current_chapter_index]
            self.log(f"Generating content for chapter: {chapter['title']}...")
            self.update_progress(0, "Starting chapter generation...")
            
            # Create the prompt for AI
            outline = self.outline_text.get(1.0, tk.END).strip()
            
            prompt = f"""Write a detailed chapter for a book with the following details:
            - Book title: {self.book_title.get()}
            - Chapter title: {chapter['title']}
            - Writing style: {self.writing_style.get()}
            
            Chapter outline:
            {outline}
            
            Please write a comprehensive chapter following this outline. The content should be detailed, well-structured, and professionally written.
            Include appropriate section headings for each major point in the outline.
            """
            
            # Send to OpenAI
            self.update_progress(20, "Generating chapter content...")
            response = openai.chat.completions.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": "You are a professional book writer with expertise in creating detailed, engaging book chapters."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=4000
            )
            
            # Get chapter content
            chapter_content = response.choices[0].message.content
            
            # Generate image/diagram if requested
            image_data = None
            if self.include_images.get() or self.include_diagrams.get():
                self.update_progress(50, "Generating visual content...")
                
                if self.include_images.get():
                    self.log("Generating image for chapter...")
                    
                    # Create a prompt for the image
                    image_prompt = f"Create an image representing the key concepts in this chapter about {chapter['title']}."
                    
                    try:
                        image_response = openai.images.generate(
                            model="dall-e-3",
                            prompt=image_prompt,
                            size="1024x1024",
                            n=1
                        )
                        
                        # Get the image URL
                        image_url = image_response.data[0].url
                        self.log(f"Image generated successfully")
                        
                        # Download the image
                        img_response = requests.get(image_url)
                        img = Image.open(io.BytesIO(img_response.content))
                        
                        # Resize for display
                        img = img.resize((400, 400), Image.LANCZOS)
                        photo = ImageTk.PhotoImage(img)
                        self.image_label.config(image=photo)
                        self.image_label.image = photo  # Keep a reference
                        
                        # Save image to include in docx
                        image_data = {
                            "type": "image",
                            "data": img_response.content
                        }
                        
                    except Exception as e:
                        self.log(f"Error generating image: {str(e)}")
                        
                elif self.include_diagrams.get():
                    self.log("Generating diagram for chapter...")
                    
                    # Generate a simple chart for demonstration
                    try:
                        fig, ax = plt.subplots(figsize=(8, 6))
                        
                        # Extract some data from the chapter to visualize
                        # This is a simple example - in reality, you'd want to analyze the text
                        # to create a meaningful chart based on the content
                        data = [5, 7, 3, 8, 6]
                        labels = ['Point 1', 'Point 2', 'Point 3', 'Point 4', 'Point 5']
                        
                        ax.bar(labels, data)
                        ax.set_title(f"Key Points in {chapter['title']}")
                        ax.set_ylabel('Importance')
                        
                        # Save to BytesIO
                        buf = io.BytesIO()
                        canvas = FigureCanvasAgg(fig)
                        canvas.print_png(buf)
                        
                        # Convert to PhotoImage for display
                        buf.seek(0)
                        img = Image.open(buf)
                        img = img.resize((500, 400), Image.LANCZOS)
                        photo = ImageTk.PhotoImage(img)
                        self.image_label.config(image=photo)
                        self.image_label.image = photo  # Keep a reference
                        
                        # Save diagram to include in docx
                        buf.seek(0)
                        image_data = {
                            "type": "diagram",
                            "data": buf.read()
                        }
                        plt.close(fig)
                        
                    except Exception as e:
                        self.log(f"Error generating diagram: {str(e)}")
            
            # Display the content in the preview area
            self.preview_text.delete(1.0, tk.END)
            self.preview_text.insert(tk.END, chapter_content)
            
            # Save the generated content for later use
            chapter["generated_content"] = chapter_content
            if image_data:
                chapter["image_data"] = image_data
            
            self.update_progress(100, "Chapter generation completed")
            self.log("Chapter content generated successfully")
            messagebox.showinfo("Success", "Chapter content generated successfully")
            
        except Exception as e:
            self.log(f"Error generating chapter content: {str(e)}")
            self.update_progress(0, "Error in content generation")
            messagebox.showerror("Error", f"Failed to generate chapter content: {str(e)}")
    
    def save_generated_chapter(self):
        if not self.chapters:
            messagebox.showerror("Error", "No chapters available.")
            return
            
        if self.current_chapter_index < 0 or self.current_chapter_index >= len(self.chapters):
            messagebox.showerror("Error", "No chapter selected.")
            return
            
        chapter = self.chapters[self.current_chapter_index]
        
        if not hasattr(chapter, "generated_content") and "generated_content" not in chapter:
            messagebox.showerror("Error", "No generated content available for this chapter.")
            return
            
        try:
            # Create output directory if it doesn't exist
            os.makedirs(self.output_dir.get(), exist_ok=True)
            
            # Create a filename
            safe_title = re.sub(r'[^\w\s-]', '', chapter["title"]).strip().replace(' ', '_')
            filepath = os.path.join(self.output_dir.get(), f"Generated_{safe_title}.docx")
            
            # Create a new document
            doc = Document()
            
            # Add title
            doc.add_heading(chapter["title"], level=1)
            
            # Add content - split by paragraphs and add each as a paragraph
            content_paragraphs = chapter["generated_content"].split("\n\n")
            for para_text in content_paragraphs:
                if not para_text.strip():
                    continue
                    
                # Check if this is a heading (starts with #, ##, etc.)
                if re.match(r'^#+\s+', para_text):
                    # Count the number of # to determine heading level
                    level = len(re.match(r'^(#+)\s+', para_text).group(1))
                    text = re.sub(r'^#+\s+', '', para_text)
                    doc.add_heading(text, level=min(level+1, 9))
                else:
                    doc.add_paragraph(para_text)
            
            # Add image if available
            if "image_data" in chapter:
                image_data = chapter["image_data"]
                if image_data:
                    img_buf = io.BytesIO(image_data["data"])
                    doc.add_picture(img_buf, width=Inches(5.5))
                    caption = doc.add_paragraph(f"Figure: Illustration for {chapter['title']}")
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption.runs[0].italic = True
            
            # Save the document
            doc.save(filepath)
            
            self.log(f"Generated chapter saved to {filepath}")
            messagebox.showinfo("Success", f"Generated chapter saved to {filepath}")
            
        except Exception as e:
            self.log(f"Error saving generated chapter: {str(e)}")
            messagebox.showerror("Error", f"Failed to save generated chapter: {str(e)}")


if __name__ == "__main__":
    root = tk.Tk()
    app = BookProcessor(root)
    root.mainloop()

