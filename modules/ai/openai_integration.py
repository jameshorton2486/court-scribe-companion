
import openai
import os
from tkinter import messagebox
import json
import time
from typing import Optional, Dict, Any, List

def get_api_key(app) -> Optional[str]:
    """Get OpenAI API key with improved handling of environment variables."""
    # First try to get API key from environment variables
    api_key = os.environ.get('OPENAI_API_KEY')
    
    # If not in environment, get from app
    if not api_key:
        api_key = app.openai_api_key.get().strip()
        if not api_key:
            app.log("No OpenAI API key found in environment or application")
            return None
    else:
        # Update the app's key field if found in environment
        if app.openai_api_key.get() != api_key:
            app.openai_api_key.set(api_key)
            app.log("Using OpenAI API key from environment variables")
    
    return api_key

def test_openai_connection(app):
    """Test connection to OpenAI API with improved error handling."""
    app.update_progress(10, "Testing OpenAI connection...")
    
    # Get API key with improved handling
    api_key = get_api_key(app)
    if not api_key:
        messagebox.showerror("Error", "No OpenAI API key found. Please set the OPENAI_API_KEY environment variable or enter it in the application.")
        app.update_progress(0, "OpenAI connection failed: No API key")
        return
    
    try:
        # Set up the client
        client = openai.OpenAI(api_key=api_key)
        
        # Simple test request
        app.log("Sending test request to OpenAI API...")
        response = client.chat.completions.create(
            model=app.openai_model.get(),
            messages=[
                {"role": "system", "content": "You are a helpful assistant."},
                {"role": "user", "content": "Hello, this is a test message."}
            ],
            max_tokens=20
        )
        
        # Log response content for debugging
        app.log(f"Response received: {response.choices[0].message.content}")
        
        app.log("OpenAI connection test successful!")
        app.update_progress(100, "OpenAI connection successful")
        messagebox.showinfo("Success", "Successfully connected to OpenAI API using " + 
                          ("environment variable" if os.environ.get('OPENAI_API_KEY') else "application input"))
        
    except openai.RateLimitError as e:
        app.log(f"OpenAI rate limit error: {str(e)}")
        app.update_progress(0, "OpenAI connection failed: Rate limit exceeded")
        messagebox.showerror("Rate Limit Error", f"OpenAI API rate limit exceeded. Please try again later: {str(e)}")
    except openai.APIError as e:
        app.log(f"OpenAI API error: {str(e)}")
        app.update_progress(0, "OpenAI connection failed: API error")
        messagebox.showerror("API Error", f"OpenAI API error: {str(e)}")
    except openai.APIConnectionError as e:
        app.log(f"OpenAI connection error: {str(e)}")
        app.update_progress(0, "OpenAI connection failed: Connection error")
        messagebox.showerror("Connection Error", f"Failed to connect to OpenAI API: {str(e)}")
    except Exception as e:
        app.log(f"OpenAI connection error: {str(e)}")
        app.update_progress(0, "OpenAI connection failed")
        messagebox.showerror("Error", f"Failed to connect to OpenAI: {str(e)}")

def enhance_text_with_ai(app, text, enhancement_type='grammar', 
                         review_grammar=True, review_coherence=True, 
                         suggest_improvements=True, intensity=0.5) -> Optional[str]:
    """
    Use OpenAI to enhance text content with improved error handling and expanded options.
    
    Args:
        app: The application instance
        text: The text to enhance
        enhancement_type: Type of enhancement ('grammar', 'expand', 'simplify', or 'style')
        review_grammar: Whether to review grammar
        review_coherence: Whether to review coherence
        suggest_improvements: Whether to suggest improvements
        intensity: Enhancement intensity (0.1 to 1.0)
    
    Returns:
        Enhanced text or None if there was an error
    """
    app.log(f"Enhancing text with AI ({enhancement_type}, intensity: {intensity})...")
    
    # Get API key with improved handling
    api_key = get_api_key(app)
    if not api_key:
        app.log("No OpenAI API key available")
        app.update_progress(0, "AI enhancement failed: No API key")
        return None
    
    # Set up the client
    client = openai.OpenAI(api_key=api_key)
    
    # Map intensity to temperature and other parameters
    temperature = min(0.7, max(0.1, 0.1 + intensity * 0.5))
    
    max_retries = 3
    retry_count = 0
    
    while retry_count < max_retries:
        try:
            # Build system prompt based on enhancement type and options
            system_prompt = "You are an expert editor and writing assistant. "
            
            if enhancement_type == 'grammar':
                system_prompt += f"Your task is to improve the text while maintaining its original meaning and style. Apply changes with {intensity * 100:.0f}% intensity - higher means more significant rewrites."
                
                if review_grammar:
                    system_prompt += " Fix any grammatical errors, spelling mistakes, and punctuation."
                    
                if review_coherence:
                    system_prompt += " Ensure the text is coherent, flows naturally, and has logical transitions."
                    
                if suggest_improvements:
                    system_prompt += " Improve clarity and readability where possible."
            
            elif enhancement_type == 'expand':
                system_prompt += f"Your task is to expand and enrich the provided content while maintaining its original style and message. Add relevant details, examples, or explanations to make the content more comprehensive. Apply expansion with {intensity * 100:.0f}% intensity - higher means more added content."
            
            elif enhancement_type == 'simplify':
                system_prompt += f"Your task is to simplify the provided content without losing its essential meaning. Make it more accessible to a general audience by reducing complexity and using simpler language. Apply simplification with {intensity * 100:.0f}% intensity - higher means more significant simplification."
                
            elif enhancement_type == 'style':
                style_type = app.style_type.get() if hasattr(app, 'style_type') else "professional"
                system_prompt += f"Your task is to rewrite the text in a {style_type} style while preserving its meaning. Apply style changes with {intensity * 100:.0f}% intensity - higher means more significant style adjustments."
            
            app.log(f"Using system prompt: {system_prompt}")
            
            # Prepare user prompt with the text to enhance
            user_prompt = f"Please enhance the following text according to the instructions. Return ONLY the improved text without additional commentary:\n\n{text}"
            
            # Call OpenAI API with selected model
            model = app.openai_model.get()
            app.log(f"Using model: {model}")
            
            # Call OpenAI API
            response = client.chat.completions.create(
                model=model,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ],
                max_tokens=4000,
                temperature=temperature
            )
            
            # Extract enhanced text from response
            enhanced_text = response.choices[0].message.content.strip()
            
            app.log("AI enhancement completed successfully")
            return enhanced_text
            
        except openai.RateLimitError:
            retry_count += 1
            app.log(f"Rate limit exceeded. Retrying in {2 ** retry_count} seconds... (Attempt {retry_count}/{max_retries})")
            app.update_progress(10 * retry_count, f"Retrying... ({retry_count}/{max_retries})")
            time.sleep(2 ** retry_count)  # Exponential backoff
            
        except (openai.APIError, openai.APIConnectionError) as e:
            app.log(f"API error in AI text enhancement: {str(e)}")
            app.update_progress(0, f"AI enhancement failed: {type(e).__name__}")
            return None
            
        except Exception as e:
            app.log(f"Error in AI text enhancement: {str(e)}")
            app.update_progress(0, "AI enhancement failed")
            return None
    
    # If we've exhausted retries
    app.log("Maximum retries exceeded for OpenAI API call")
    app.update_progress(0, "AI enhancement failed: Max retries exceeded")
    return None

def generate_ai_toc(app):
    """Generate a table of contents using AI with improved error handling."""
    if not hasattr(app, 'chapters') or not app.chapters:
        messagebox.showerror("Error", "No chapters available for TOC generation")
        return
    
    app.log("Generating AI-assisted table of contents...")
    app.update_progress(0, "Starting AI TOC generation...")
    
    # Get API key with improved handling
    api_key = get_api_key(app)
    if not api_key:
        messagebox.showerror("Error", "No OpenAI API key available")
        app.log("No OpenAI API key available")
        app.update_progress(0, "AI TOC generation failed: No API key")
        return
    
    # Set up the client
    client = openai.OpenAI(api_key=api_key)
    
    try:
        # Prepare chapter titles and excerpts for TOC generation
        toc_input = []
        
        for i, chapter in enumerate(app.chapters):
            # Get chapter title
            title = chapter['title']
            
            # Get a brief excerpt (first 200 chars) from the chapter content
            excerpt = ""
            for para in chapter['content'][1:5]:  # Skip title para, take next few
                excerpt += para.text + " "
                if len(excerpt) > 200:
                    excerpt = excerpt[:200] + "..."
                    break
            
            toc_input.append(f"Chapter {i+1}: {title}\nExcerpt: {excerpt}\n")
        
        # Combine all chapter info
        all_chapters_info = "\n".join(toc_input)
        
        # Call OpenAI API to generate TOC
        system_prompt = "You are an expert book editor specialized in creating detailed tables of contents. Analyze the chapter titles and excerpts to create a hierarchical table of contents with main sections and subsections."
        
        user_prompt = f"Based on these chapter titles and excerpts, please generate a detailed table of contents with appropriate hierarchy for a book titled '{app.book_title.get()}':\n\n{all_chapters_info}\n\nFormat your response as a JSON array with each entry having 'title', 'level' (1 for main entries, 2 for sub-entries, 3 for sub-sub-entries), and 'chapter_index' properties."
        
        app.log("Sending TOC generation request to OpenAI...")
        model = app.openai_model.get()
        app.log(f"Using model: {model}")
        
        response = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            max_tokens=2000,
            temperature=0.3
        )
        
        # Extract TOC from response
        toc_response = response.choices[0].message.content.strip()
        
        # Try to parse the JSON response
        try:
            # Look for JSON array in the response
            json_start = toc_response.find('[')
            json_end = toc_response.rfind(']') + 1
            
            if json_start >= 0 and json_end > json_start:
                json_str = toc_response[json_start:json_end]
                ai_toc = json.loads(json_str)
                
                # Update the app's TOC
                app.toc = []
                for item in ai_toc:
                    toc_item = {
                        'title': item['title'],
                        'level': item['level'],
                        'index': item.get('chapter_index', 0)
                    }
                    app.toc.append(toc_item)
                
                app.log(f"AI-generated TOC with {len(app.toc)} entries")
                app.update_progress(100, "AI TOC generation completed")
                messagebox.showinfo("Success", "AI-assisted table of contents generated successfully")
                return True
            else:
                app.log("Could not find valid JSON in API response")
                app.update_progress(0, "AI TOC generation failed: Invalid response format")
                messagebox.showerror("Error", "Failed to generate TOC: invalid response format")
                return False
                
        except json.JSONDecodeError as e:
            app.log(f"Error parsing TOC JSON: {str(e)}")
            app.log(f"Raw response: {toc_response}")
            app.update_progress(0, "AI TOC generation failed: JSON parse error")
            messagebox.showerror("Error", "Failed to parse AI-generated TOC")
            return False
        
    except openai.RateLimitError as e:
        app.log(f"OpenAI rate limit error: {str(e)}")
        app.update_progress(0, "AI TOC generation failed: Rate limit exceeded")
        messagebox.showerror("Rate Limit Error", f"OpenAI API rate limit exceeded. Please try again later: {str(e)}")
        return False
    except (openai.APIError, openai.APIConnectionError) as e:
        app.log(f"API error in TOC generation: {str(e)}")
        app.update_progress(0, f"AI TOC generation failed: {type(e).__name__}")
        messagebox.showerror("API Error", f"OpenAI API error: {str(e)}")
        return False
    except Exception as e:
        app.log(f"Error in AI TOC generation: {str(e)}")
        app.update_progress(0, "AI TOC generation failed")
        messagebox.showerror("Error", f"Failed to generate AI TOC: {str(e)}")
        return False

def review_with_ai(app):
    """Review document content with AI to provide feedback and suggestions with improved error handling."""
    if not hasattr(app, 'chapters') or not app.chapters:
        messagebox.showerror("Error", "No chapters available for review")
        return
    
    app.log("Starting AI content review...")
    app.update_progress(0, "Starting AI content review...")
    
    # Get API key with improved handling
    api_key = get_api_key(app)
    if not api_key:
        messagebox.showerror("Error", "No OpenAI API key available")
        app.log("No OpenAI API key available")
        app.update_progress(0, "AI review failed: No API key")
        return
    
    # Set up the client
    client = openai.OpenAI(api_key=api_key)
    
    try:
        # Analyze each chapter and compile feedback
        all_feedback = []
        total_chapters = len(app.chapters)
        
        for i, chapter in enumerate(app.chapters):
            app.update_progress((i / total_chapters) * 100, f"Reviewing chapter {i+1} of {total_chapters}")
            
            # Get chapter content
            chapter_text = "\n".join([p.text for p in chapter['content']])
            
            # Prepare review parameters based on user selections
            review_params = []
            if app.review_grammar.get():
                review_params.append("grammar and style")
            if app.review_coherence.get():
                review_params.append("content coherence and flow")
            if app.suggest_improvements.get():
                review_params.append("content improvement suggestions")
            
            review_focus = ", ".join(review_params)
            
            # Skip if no review parameters selected
            if not review_params:
                app.log(f"Skipping chapter {i+1} (no review parameters selected)")
                continue
            
            # Call OpenAI API for review
            system_prompt = "You are an expert editor and writing coach. Provide a concise, helpful review of the text focusing on the requested aspects. Be specific and constructive with your feedback."
            
            user_prompt = f"Please review this text and provide feedback on {review_focus}. The text is chapter {i+1} ('{chapter['title']}') of the book '{app.book_title.get()}':\n\n{chapter_text}\n\nFormat your response as a JSON object with sections for each requested review aspect, with 'issues' (array of specific issues) and 'suggestions' (array of improvement ideas) for each section."
            
            app.log(f"Sending review request for chapter {i+1} to OpenAI...")
            model = app.openai_model.get()
            
            try:
                response = client.chat.completions.create(
                    model=model,
                    messages=[
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": user_prompt}
                    ],
                    max_tokens=2000,
                    temperature=0.3
                )
                
                # Process review feedback
                feedback = response.choices[0].message.content.strip()
                all_feedback.append(f"## Chapter {i+1}: {chapter['title']}\n\n{feedback}\n\n")
            except openai.RateLimitError as e:
                app.log(f"Rate limit exceeded for chapter {i+1}: {str(e)}")
                all_feedback.append(f"## Chapter {i+1}: {chapter['title']}\n\n*Error: API rate limit exceeded for this chapter.*\n\n")
                # Continue with next chapter instead of stopping the whole process
                time.sleep(5)  # Wait a bit before next request
                continue
        
        # Combine all feedback
        combined_feedback = "\n".join(all_feedback)
        
        # Create a report document
        from docx import Document
        report_doc = Document()
        
        # Add title
        report_doc.add_heading(f"AI Review Report: {app.book_title.get()}", level=1)
        
        # Add review content
        for line in combined_feedback.split('\n'):
            # Check if line is a heading
            if line.startswith('##'):
                level = line.count('#')
                report_doc.add_heading(line.lstrip('#').strip(), level=level)
            else:
                report_doc.add_paragraph(line)
        
        # Save the report
        report_filename = f"{app.book_title.get().replace(' ', '_')}_AI_Review.docx"
        
        # Ensure output directory exists
        import os
        os.makedirs(app.output_dir.get(), exist_ok=True)
        
        report_path = os.path.join(app.output_dir.get(), report_filename)
        report_doc.save(report_path)
        
        app.update_progress(100, "AI review completed")
        app.log(f"AI review completed and saved to: {report_path}")
        messagebox.showinfo("Success", f"AI review completed and saved to: {report_path}")
        
    except openai.RateLimitError as e:
        app.log(f"OpenAI rate limit error: {str(e)}")
        app.update_progress(0, "AI review failed: Rate limit exceeded")
        messagebox.showerror("Rate Limit Error", f"OpenAI API rate limit exceeded. Please try again later: {str(e)}")
    except (openai.APIError, openai.APIConnectionError) as e:
        app.log(f"API error in AI review: {str(e)}")
        app.update_progress(0, f"AI review failed: {type(e).__name__}")
        messagebox.showerror("API Error", f"OpenAI API error: {str(e)}")
    except Exception as e:
        app.log(f"Error in AI review: {str(e)}")
        app.update_progress(0, "AI review failed")
        messagebox.showerror("Error", f"Failed to complete AI review: {str(e)}")
