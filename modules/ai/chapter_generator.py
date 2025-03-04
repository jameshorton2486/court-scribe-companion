
import threading
import io
import os
import re
import tkinter as tk
import openai
import requests
from PIL import Image, ImageTk
import matplotlib.pyplot as plt
from matplotlib.backends.backend_agg import FigureCanvasAgg
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from tkinter import messagebox

def generate_chapter_content(app):
    if not app.chapters:
        messagebox.showerror("Error", "No chapters available. Please process a document first.")
        return
    
    if app.current_chapter_index < 0 or app.current_chapter_index >= len(app.chapters):
        messagebox.showerror("Error", "Please select a valid chapter first.")
        return
    
    outline = app.outline_text.get(1.0, tk.END).strip()
    if not outline:
        messagebox.showerror("Error", "Please provide an outline for the chapter.")
        return
    
    api_key = app.openai_api_key.get()
    if not api_key:
        messagebox.showerror("Error", "Please enter an OpenAI API key")
        return
    
    openai.api_key = api_key
    
    # Start chapter generation in a separate thread
    threading.Thread(target=_generate_chapter_thread, args=(app,), daemon=True).start()

def _generate_chapter_thread(app):
    try:
        chapter = app.chapters[app.current_chapter_index]
        app.log(f"Generating content for chapter: {chapter['title']}...")
        app.update_progress(0, "Starting chapter generation...")
        
        # Create the prompt for AI
        outline = app.outline_text.get(1.0, tk.END).strip()
        
        prompt = f"""Write a detailed chapter for a book with the following details:
        - Book title: {app.book_title.get()}
        - Chapter title: {chapter['title']}
        - Writing style: {app.writing_style.get()}
        
        Chapter outline:
        {outline}
        
        Please write a comprehensive chapter following this outline. The content should be detailed, well-structured, and professionally written.
        Include appropriate section headings for each major point in the outline.
        """
        
        # Send to OpenAI
        app.update_progress(20, "Generating chapter content...")
        response = openai.chat.completions.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": "You are a professional book writer with expertise in creating detailed, engaging book chapters."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=4000
        )
        
        # Get chapter content
        chapter_content = response.choices[0].message.content
        
        # Generate image/diagram if requested
        image_data = None
        if app.include_images.get() or app.include_diagrams.get():
            app.update_progress(50, "Generating visual content...")
            
            if app.include_images.get():
                app.log("Generating image for chapter...")
                
                # Create a prompt for the image
                image_prompt = f"Create an image representing the key concepts in this chapter about {chapter['title']}."
                
                try:
                    image_response = openai.images.generate(
                        model="dall-e-3",
                        prompt=image_prompt,
                        size="1024x1024",
                        n=1
                    )
                    
                    # Get the image URL
                    image_url = image_response.data[0].url
                    app.log(f"Image generated successfully")
                    
                    # Download the image
                    img_response = requests.get(image_url)
                    img = Image.open(io.BytesIO(img_response.content))
                    
                    # Resize for display
                    img = img.resize((400, 400), Image.LANCZOS)
                    photo = ImageTk.PhotoImage(img)
                    app.image_label.config(image=photo)
                    app.image_label.image = photo  # Keep a reference
                    
                    # Save image to include in docx
                    image_data = {
                        "type": "image",
                        "data": img_response.content
                    }
                    
                except Exception as e:
                    app.log(f"Error generating image: {str(e)}")
                    
            elif app.include_diagrams.get():
                app.log("Generating diagram for chapter...")
                
                # Generate a simple chart for demonstration
                try:
                    fig, ax = plt.subplots(figsize=(8, 6))
                    
                    # Extract some data from the chapter to visualize
                    # This is a simple example - in reality, you'd want to analyze the text
                    # to create a meaningful chart based on the content
                    data = [5, 7, 3, 8, 6]
                    labels = ['Point 1', 'Point 2', 'Point 3', 'Point 4', 'Point 5']
                    
                    ax.bar(labels, data)
                    ax.set_title(f"Key Points in {chapter['title']}")
                    ax.set_ylabel('Importance')
                    
                    # Save to BytesIO
                    buf = io.BytesIO()
                    canvas = FigureCanvasAgg(fig)
                    canvas.print_png(buf)
                    
                    # Convert to PhotoImage for display
                    buf.seek(0)
                    img = Image.open(buf)
                    img = img.resize((500, 400), Image.LANCZOS)
                    photo = ImageTk.PhotoImage(img)
                    app.image_label.config(image=photo)
                    app.image_label.image = photo  # Keep a reference
                    
                    # Save diagram to include in docx
                    buf.seek(0)
                    image_data = {
                        "type": "diagram",
                        "data": buf.read()
                    }
                    plt.close(fig)
                    
                except Exception as e:
                    app.log(f"Error generating diagram: {str(e)}")
        
        # Display the content in the preview area
        app.preview_text.delete(1.0, tk.END)
        app.preview_text.insert(tk.END, chapter_content)
        
        # Save the generated content for later use
        chapter["generated_content"] = chapter_content
        if image_data:
            chapter["image_data"] = image_data
        
        app.update_progress(100, "Chapter generation completed")
        app.log("Chapter content generated successfully")
        messagebox.showinfo("Success", "Chapter content generated successfully")
        
    except Exception as e:
        app.log(f"Error generating chapter content: {str(e)}")
        app.update_progress(0, "Error in content generation")
        messagebox.showerror("Error", f"Failed to generate chapter content: {str(e)}")

def save_generated_chapter(app):
    if not app.chapters:
        messagebox.showerror("Error", "No chapters available.")
        return
        
    if app.current_chapter_index < 0 or app.current_chapter_index >= len(app.chapters):
        messagebox.showerror("Error", "No chapter selected.")
        return
        
    chapter = app.chapters[app.current_chapter_index]
    
    if not hasattr(chapter, "generated_content") and "generated_content" not in chapter:
        messagebox.showerror("Error", "No generated content available for this chapter.")
        return
        
    try:
        # Create output directory if it doesn't exist
        os.makedirs(app.output_dir.get(), exist_ok=True)
        
        # Create a filename
        safe_title = re.sub(r'[^\w\s-]', '', chapter["title"]).strip().replace(' ', '_')
        filepath = os.path.join(app.output_dir.get(), f"Generated_{safe_title}.docx")
        
        # Create a new document
        doc = Document()
        
        # Add title
        doc.add_heading(chapter["title"], level=1)
        
        # Add content - split by paragraphs and add each as a paragraph
        content_paragraphs = chapter["generated_content"].split("\n\n")
        for para_text in content_paragraphs:
            if not para_text.strip():
                continue
                
            # Check if this is a heading (starts with #, ##, etc.)
            if re.match(r'^#+\s+', para_text):
                # Count the number of # to determine heading level
                level = len(re.match(r'^(#+)\s+', para_text).group(1))
                text = re.sub(r'^#+\s+', '', para_text)
                doc.add_heading(text, level=min(level+1, 9))
            else:
                doc.add_paragraph(para_text)
        
        # Add image if available
        if "image_data" in chapter:
            image_data = chapter["image_data"]
            if image_data:
                img_buf = io.BytesIO(image_data["data"])
                doc.add_picture(img_buf, width=Inches(5.5))
                caption = doc.add_paragraph(f"Figure: Illustration for {chapter['title']}")
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption.runs[0].italic = True
        
        # Save the document
        doc.save(filepath)
        
        app.log(f"Generated chapter saved to {filepath}")
        messagebox.showinfo("Success", f"Generated chapter saved to {filepath}")
        
    except Exception as e:
        app.log(f"Error saving generated chapter: {str(e)}")
        messagebox.showerror("Error", f"Failed to save generated chapter: {str(e)}")
