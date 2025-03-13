
"""
AI-powered Table of Contents Generator

This module provides functionality for generating an enhanced table of contents
using AI (OpenAI GPT-4) to create more detailed section structure.
"""

import threading
import time
import os
import tkinter as tk
import openai
from docx import Document
from tkinter import messagebox
from modules.utils.error_handler import ErrorHandler
from modules.ai.openai.api_client import get_api_key, execute_with_retry

def generate_ai_toc(app):
    """
    Generate an AI-assisted table of contents for the processed document.
    
    Args:
        app: The application instance
    """
    if not app.chapters:
        messagebox.showerror("Error", "No chapters found. Please process a document first.")
        return
    
    # First try to get API key from environment variables
    api_key = os.environ.get('OPENAI_API_KEY')
    
    # If not in environment, use the one from the application
    if not api_key:
        api_key = app.openai_api_key.get()
        if not api_key:
            messagebox.showerror("Error", "No OpenAI API key found. Please set the OPENAI_API_KEY environment variable or enter it in the application.")
            return
    
    # Start TOC generation in a separate thread
    threading.Thread(target=_generate_ai_toc_thread, args=(app, api_key), daemon=True).start()

def _generate_ai_toc_thread(app, api_key):
    """
    Thread function for generating AI-assisted table of contents.
    
    Args:
        app: The application instance
        api_key: OpenAI API key to use
    """
    try:
        ErrorHandler.log_operation_start("AI TOC generation")
        start_time = time.time()
        
        app.log("Generating AI-assisted table of contents...")
        app.update_progress(0, "Starting TOC generation...")
        
        # Setup OpenAI client
        client = openai.OpenAI(api_key=api_key)
        
        # Extract chapter titles and brief content
        chapters_info = []
        app.update_progress(10, "Analyzing chapters...")
        
        for chapter in app.chapters:
            # Get first 3 paragraphs of content (or fewer if there aren't 3)
            preview = ""
            for i, para in enumerate(chapter["content"][:3]):
                preview += para.text + "\n"
            
            chapters_info.append({
                "title": chapter["title"],
                "preview": preview
            })
        
        # Create prompt for OpenAI
        chapters_text = ""
        for i, ch in enumerate(chapters_info):
            chapters_text += f"Chapter {i+1}: {ch['title']}\nPreview: {ch['preview']}\n\n"
        
        prompt = f"""Based on the following chapter titles and content previews, generate a comprehensive table of contents for the book titled "{app.book_title.get()}".
        For each chapter, add 3-5 detailed subsections that would logically organize the content.
        Format the output as follows:
        
        Chapter 1: [Chapter Title]
        1.1 [Subsection title]
        1.2 [Subsection title]
        ...and so on
        
        Chapters:
        {chapters_text}
        """
        
        # Send to OpenAI
        app.update_progress(30, "Querying OpenAI...")
        app.log("Sending request to OpenAI for TOC generation...")
        
        try:
            # Use the execute_with_retry function to handle rate limits
            response = execute_with_retry(
                app, 
                client.chat.completions.create,
                model="gpt-4",
                messages=[
                    {"role": "system", "content": "You are a professional book editor with expertise in organizing book content."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=2000
            )
            
            # Get response
            toc_content = response.choices[0].message.content
            
        except openai.RateLimitError as e:
            app.log(f"OpenAI rate limit exceeded: {str(e)}")
            app.update_progress(0, "TOC generation failed: Rate limit exceeded")
            
            if app.fallback_to_local.get():
                app.log("Falling back to local TOC generation...")
                app.update_progress(40, "Falling back to local TOC generation...")
                
                # Use a simpler local TOC generation approach
                from modules.document.toc_generator import generate_table_of_contents
                generate_table_of_contents(app)
                
                app.update_progress(100, "Local TOC generation completed")
                app.log("Local TOC generation completed (fallback)")
                messagebox.showinfo("Fallback Success", "Local TOC generation completed as a fallback to OpenAI.")
                return
            else:
                messagebox.showerror("Error", f"TOC generation failed: OpenAI rate limit exceeded. Try again later.")
                return
                
        except Exception as e:
            app.log(f"Error during OpenAI TOC generation: {str(e)}")
            app.update_progress(0, "TOC generation failed")
            
            # If fallback is enabled, use local generation
            if hasattr(app, 'fallback_to_local') and app.fallback_to_local.get():
                app.log("Falling back to local TOC generation...")
                app.update_progress(40, "Falling back to local TOC generation...")
                
                # Use a simpler local TOC generation approach
                from modules.document.toc_generator import generate_table_of_contents
                generate_table_of_contents(app)
                
                app.update_progress(100, "Local TOC generation completed")
                app.log("Local TOC generation completed (fallback)")
                messagebox.showinfo("Fallback Success", "Local TOC generation completed as a fallback to OpenAI.")
                return
            else:
                messagebox.showerror("Error", f"TOC generation failed: {str(e)}")
                return
        
        # Create TOC document
        app.update_progress(70, "Creating TOC document...")
        
        toc_doc = Document()
        toc_doc.add_heading(f"Table of Contents: {app.book_title.get()}", level=1)
        toc_doc.add_paragraph(f"Author: {app.author_name.get()}")
        toc_doc.add_paragraph(f"Generated on: {time.strftime('%Y-%m-%d %H:%M:%S')}")
        toc_doc.add_paragraph()
        
        # Add the TOC content
        for line in toc_content.split('\n'):
            if line.strip():
                # Format based on if it's a chapter or subsection
                if line.startswith('Chapter'):
                    p = toc_doc.add_paragraph()
                    p.add_run(line).bold = True
                else:
                    toc_doc.add_paragraph(line, style='List Bullet')
        
        # Save TOC document
        try:
            os.makedirs(app.output_dir.get(), exist_ok=True)
            toc_filepath = os.path.join(app.output_dir.get(), "AI_Table_of_Contents.docx")
            toc_doc.save(toc_filepath)
            app.log(f"TOC document saved to: {toc_filepath}")
        except Exception as e:
            app.log(f"Error saving TOC document: {str(e)}")
            messagebox.showerror("Error", f"Failed to save TOC document: {str(e)}")
            # Continue with processing even if saving fails
        
        # Also update application's TOC
        app.toc = []
        current_chapter = None
        
        for line in toc_content.split('\n'):
            if line.startswith('Chapter'):
                # This is a main chapter
                chapter_parts = line.split(':', 1)
                if len(chapter_parts) > 1:
                    chapter_number = chapter_parts[0].replace('Chapter', '').strip()
                    chapter_title = chapter_parts[1].strip()
                    
                    try:
                        idx = int(chapter_number) - 1
                        if 0 <= idx < len(app.chapters):
                            current_chapter = idx
                            app.toc.append({
                                "title": chapter_title,
                                "level": 1,
                                "index": idx
                            })
                    except ValueError:
                        pass
                    
            elif current_chapter is not None and line.strip():
                # This is a subsection
                app.toc.append({
                    "title": line.strip(),
                    "level": 2,
                    "index": current_chapter
                })
        
        # Populate chapter listbox
        app.chapter_listbox.delete(0, tk.END)
        for i, chapter in enumerate(app.chapters):
            app.chapter_listbox.insert(tk.END, f"Chapter {i+1}: {chapter['title']}")
        
        # Record timing and completion
        processing_time = time.time() - start_time
        
        app.update_progress(100, "TOC generation completed")
        app.log(f"AI TOC generation completed in {processing_time:.2f} seconds. TOC document saved to {toc_filepath}")
        
        ErrorHandler.log_operation_complete(
            "AI TOC generation", 
            success=True,
            duration=processing_time,
            details=f"Generated for {len(app.chapters)} chapters"
        )
        
        messagebox.showinfo("Success", f"AI TOC generation completed. Document saved to {toc_filepath}")
        
    except Exception as e:
        app.log(f"Error during TOC generation: {str(e)}")
        app.update_progress(0, "Error during TOC generation")
        
        # Log operation failure
        ErrorHandler.log_operation_complete(
            "AI TOC generation", 
            success=False,
            details=str(e)
        )
        
        messagebox.showerror("Error", f"Failed to generate TOC: {str(e)}")
