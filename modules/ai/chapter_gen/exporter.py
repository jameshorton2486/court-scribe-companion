
"""
Chapter Export Module

This module handles the export of AI-generated chapters to various formats,
primarily focusing on Microsoft Word (DOCX) output.
"""

import os
import re
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
from tkinter import messagebox


def save_generated_chapter(app):
    """
    Save a generated chapter to a Word document file.
    
    This function takes the currently selected chapter from the application,
    formats it appropriately with headings, paragraphs, and any associated
    images, and saves it as a Word document in the specified output directory.
    
    Args:
        app: The application instance containing chapter data and UI elements
        
    Returns:
        None
        
    Raises:
        Various exceptions related to file access or document creation issues
    """
    if not app.chapters:
        messagebox.showerror("Error", "No chapters available.")
        return
        
    if app.current_chapter_index < 0 or app.current_chapter_index >= len(app.chapters):
        messagebox.showerror("Error", "No chapter selected.")
        return
        
    chapter = app.chapters[app.current_chapter_index]
    
    if not hasattr(chapter, "generated_content") and "generated_content" not in chapter:
        messagebox.showerror("Error", "No generated content available for this chapter.")
        return
        
    try:
        # Create output directory if it doesn't exist
        os.makedirs(app.output_dir.get(), exist_ok=True)
        
        # Create a filename
        safe_title = re.sub(r'[^\w\s-]', '', chapter["title"]).strip().replace(' ', '_')
        filepath = os.path.join(app.output_dir.get(), f"Generated_{safe_title}.docx")
        
        # Create a new document
        doc = Document()
        
        # Add title
        doc.add_heading(chapter["title"], level=1)
        
        # Add content - split by paragraphs and add each as a paragraph
        content_paragraphs = chapter["generated_content"].split("\n\n")
        for para_text in content_paragraphs:
            if not para_text.strip():
                continue
                
            # Check if this is a heading (starts with #, ##, etc.)
            if re.match(r'^#+\s+', para_text):
                # Count the number of # to determine heading level
                level = len(re.match(r'^(#+)\s+', para_text).group(1))
                text = re.sub(r'^#+\s+', '', para_text)
                doc.add_heading(text, level=min(level+1, 9))
            else:
                doc.add_paragraph(para_text)
        
        # Add image if available
        if "image_data" in chapter:
            image_data = chapter["image_data"]
            if image_data:
                img_buf = io.BytesIO(image_data["data"])
                doc.add_picture(img_buf, width=Inches(5.5))
                caption = doc.add_paragraph(f"Figure: Illustration for {chapter['title']}")
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption.runs[0].italic = True
        
        # Save the document
        doc.save(filepath)
        
        app.log(f"Generated chapter saved to {filepath}")
        messagebox.showinfo("Success", f"Generated chapter saved to {filepath}")
        
    except Exception as e:
        app.log(f"Error saving generated chapter: {str(e)}")
        messagebox.showerror("Error", f"Failed to save generated chapter: {str(e)}")
