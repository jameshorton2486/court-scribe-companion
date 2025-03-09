
import os
import time
import openai
from docx import Document
from tkinter import messagebox
from modules.ai.openai.api_client import get_api_key

def review_with_ai(app):
    """Review document content with AI to provide feedback and suggestions."""
    if not _validate_document(app):
        return
    
    app.log("Starting AI content review...")
    app.update_progress(0, "Starting AI content review...")
    
    api_key = _get_api_key_with_validation(app)
    if not api_key:
        return
    
    client = openai.OpenAI(api_key=api_key)
    
    try:
        all_feedback = _process_chapters(app, client)
        _create_and_save_report(app, all_feedback)
        
    except openai.RateLimitError as e:
        _handle_rate_limit_error(app, e)
    except (openai.APIError, openai.APIConnectionError) as e:
        _handle_api_error(app, e)
    except Exception as e:
        _handle_general_error(app, e)

def _validate_document(app):
    """Validate that chapters are available for review."""
    if not hasattr(app, 'chapters') or not app.chapters:
        messagebox.showerror("Error", "No chapters available for review")
        return False
    return True

def _get_api_key_with_validation(app):
    """Get and validate OpenAI API key."""
    api_key = get_api_key(app)
    if not api_key:
        messagebox.showerror("Error", "No OpenAI API key available")
        app.log("No OpenAI API key available")
        app.update_progress(0, "AI review failed: No API key")
        return None
    return api_key

def _get_review_parameters(app):
    """Get review parameters based on user selections."""
    review_params = []
    if app.review_grammar.get():
        review_params.append("grammar and style")
    if app.review_coherence.get():
        review_params.append("content coherence and flow")
    if app.suggest_improvements.get():
        review_params.append("content improvement suggestions")
    
    return review_params

def _create_review_prompt(chapter, chapter_index, review_focus, book_title):
    """Create the system and user prompts for the review."""
    chapter_text = "\n".join([p.text for p in chapter['content']])
    
    system_prompt = "You are an expert editor and writing coach. Provide a concise, helpful review of the text focusing on the requested aspects. Be specific and constructive with your feedback."
    
    user_prompt = f"Please review this text and provide feedback on {review_focus}. The text is chapter {chapter_index+1} ('{chapter['title']}') of the book '{book_title}':\n\n{chapter_text}\n\nFormat your response as a JSON object with sections for each requested review aspect, with 'issues' (array of specific issues) and 'suggestions' (array of improvement ideas) for each section."
    
    return system_prompt, user_prompt, chapter_text

def _review_chapter(client, app, chapter, chapter_index, review_params):
    """Review a single chapter and return the feedback."""
    if not review_params:
        app.log(f"Skipping chapter {chapter_index+1} (no review parameters selected)")
        return None
    
    review_focus = ", ".join(review_params)
    system_prompt, user_prompt, _ = _create_review_prompt(
        chapter, 
        chapter_index, 
        review_focus, 
        app.book_title.get()
    )
    
    app.log(f"Sending review request for chapter {chapter_index+1} to OpenAI...")
    model = app.openai_model.get()
    
    try:
        response = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            max_tokens=2000,
            temperature=0.3
        )
        
        feedback = response.choices[0].message.content.strip()
        return f"## Chapter {chapter_index+1}: {chapter['title']}\n\n{feedback}\n\n"
    
    except openai.RateLimitError as e:
        app.log(f"Rate limit exceeded for chapter {chapter_index+1}: {str(e)}")
        return f"## Chapter {chapter_index+1}: {chapter['title']}\n\n*Error: API rate limit exceeded for this chapter.*\n\n"

def _process_chapters(app, client):
    """Process all chapters and return combined feedback."""
    all_feedback = []
    total_chapters = len(app.chapters)
    
    for i, chapter in enumerate(app.chapters):
        app.update_progress((i / total_chapters) * 100, f"Reviewing chapter {i+1} of {total_chapters}")
        
        review_params = _get_review_parameters(app)
        feedback = _review_chapter(client, app, chapter, i, review_params)
        
        if feedback:
            all_feedback.append(feedback)
            
        if i < total_chapters - 1:  # Don't wait after the last chapter
            time.sleep(1)  # Brief pause between API calls
    
    return all_feedback

def _create_and_save_report(app, all_feedback):
    """Create and save the final report document."""
    if not all_feedback:
        app.log("No feedback generated. Review complete.")
        app.update_progress(0, "AI review completed: No feedback")
        messagebox.showinfo("Notice", "Review completed, but no feedback was generated.")
        return
    
    combined_feedback = "\n".join(all_feedback)
    report_doc = Document()
    
    # Add title
    report_doc.add_heading(f"AI Review Report: {app.book_title.get()}", level=1)
    
    # Add review content
    for line in combined_feedback.split('\n'):
        if line.startswith('##'):
            level = line.count('#')
            report_doc.add_heading(line.lstrip('#').strip(), level=level)
        else:
            report_doc.add_paragraph(line)
    
    # Save the report
    report_filename = f"{app.book_title.get().replace(' ', '_')}_AI_Review.docx"
    
    # Ensure output directory exists
    os.makedirs(app.output_dir.get(), exist_ok=True)
    
    report_path = os.path.join(app.output_dir.get(), report_filename)
    report_doc.save(report_path)
    
    app.update_progress(100, "AI review completed")
    app.log(f"AI review completed and saved to: {report_path}")
    messagebox.showinfo("Success", f"AI review completed and saved to: {report_path}")

def _handle_rate_limit_error(app, error):
    """Handle OpenAI rate limit errors."""
    app.log(f"OpenAI rate limit error: {str(error)}")
    app.update_progress(0, "AI review failed: Rate limit exceeded")
    messagebox.showerror("Rate Limit Error", f"OpenAI API rate limit exceeded. Please try again later: {str(error)}")

def _handle_api_error(app, error):
    """Handle OpenAI API errors."""
    app.log(f"API error in AI review: {str(error)}")
    app.update_progress(0, f"AI review failed: {type(error).__name__}")
    messagebox.showerror("API Error", f"OpenAI API error: {str(error)}")

def _handle_general_error(app, error):
    """Handle general exceptions."""
    app.log(f"Error in AI review: {str(error)}")
    app.update_progress(0, "AI review failed")
    messagebox.showerror("Error", f"Failed to complete AI review: {str(error)}")
