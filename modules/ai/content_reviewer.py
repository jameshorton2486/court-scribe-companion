
"""
Content Reviewer Module

This module provides functionality for AI-assisted review of document content,
providing feedback on quality, style, and structure.
"""

import threading
import time
import os
import traceback
from docx import Document
from tkinter import messagebox
from modules.utils.error_handler import ErrorHandler
from modules.ai.openai.content_reviewer import review_with_ai as openai_review

def review_with_ai(app):
    """
    Wrapper function for AI content review functionality.
    Delegates to the OpenAI implementation in a separate thread.
    
    Args:
        app: The application instance
    """
    if not app.chapters:
        messagebox.showerror("Error", "No chapters found. Please process a document first.")
        return
    
    # Start AI review in a separate thread
    threading.Thread(target=_review_with_ai_thread, args=(app,), daemon=True).start()

def _review_with_ai_thread(app):
    """
    Thread function to perform AI review.
    
    Args:
        app: The application instance
    """
    try:
        ErrorHandler.log_operation_start("AI content review")
        start_time = time.time()
        
        app.log("Starting AI content review...")
        app.update_progress(0, "Starting AI review...")
        
        # Call the implementation from the openai module
        result = openai_review(app)
        
        # Calculate timing
        processing_time = time.time() - start_time
        
        # Log completion
        ErrorHandler.log_operation_complete(
            "AI content review",
            success=True, 
            duration=processing_time,
            details=f"Reviewed {len(app.chapters)} chapters"
        )
        
        return result
        
    except Exception as e:
        error_details = traceback.format_exc()
        app.log(f"Error during AI review: {str(e)}")
        app.log.error(f"Error details: {error_details}")
        app.update_progress(0, "Error during AI review")
        
        # Check if we should fall back to a local review method
        if hasattr(app, 'fallback_to_local') and app.fallback_to_local.get():
            app.log("Attempting to fall back to local review...")
            try:
                # Simple local review implementation
                _perform_local_review(app)
                
                # Log fallback success
                ErrorHandler.log_operation_complete(
                    "AI content review (fallback)",
                    success=True,
                    details="Used local fallback method"
                )
                
                return True
                
            except Exception as fallback_error:
                app.log(f"Fallback review also failed: {str(fallback_error)}")
                ErrorHandler.log_operation_complete(
                    "AI content review",
                    success=False,
                    details=f"Original error: {str(e)}. Fallback also failed: {str(fallback_error)}"
                )
                
        else:
            # Log operation failure
            ErrorHandler.log_operation_complete(
                "AI content review",
                success=False,
                details=str(e)
            )
        
        messagebox.showerror("Error", f"Failed to complete AI review: {str(e)}")
        return False

def _perform_local_review(app):
    """
    Perform a simple local review without AI services.
    This is used as a fallback if the AI review fails.
    
    Args:
        app: The application instance
    """
    app.log("Performing local review (fallback method)...")
    app.update_progress(30, "Performing local review...")
    
    # Create a simple review document
    review_doc = Document()
    review_doc.add_heading(f"Document Review Report: {app.book_title.get()}", level=1)
    review_doc.add_paragraph(f"Author: {app.author_name.get()}")
    review_doc.add_paragraph(f"Generated on: {time.strftime('%Y-%m-%d %H:%M:%S')}")
    review_doc.add_paragraph(f"Note: This is a basic fallback review generated locally.")
    
    # Add document statistics
    review_doc.add_heading("Document Statistics", level=2)
    
    total_paragraphs = 0
    total_words = 0
    chapter_stats = []
    
    for i, chapter in enumerate(app.chapters):
        chapter_paragraphs = len(chapter['content'])
        chapter_words = sum(len(p.text.split()) for p in chapter['content'])
        
        total_paragraphs += chapter_paragraphs
        total_words += chapter_words
        
        chapter_stats.append({
            'title': chapter['title'],
            'paragraphs': chapter_paragraphs,
            'words': chapter_words
        })
    
    stats_para = review_doc.add_paragraph()
    stats_para.add_run(f"Total chapters: {len(app.chapters)}").bold = True
    stats_para.add_run(f"\nTotal paragraphs: {total_paragraphs}")
    stats_para.add_run(f"\nTotal words: {total_words}")
    stats_para.add_run(f"\nAverage words per chapter: {total_words // max(1, len(app.chapters))}")
    
    # Add chapter breakdowns
    review_doc.add_heading("Chapter Breakdown", level=2)
    
    for i, stats in enumerate(chapter_stats):
        chapter_heading = review_doc.add_paragraph()
        chapter_heading.add_run(f"Chapter {i+1}: {stats['title']}").bold = True
        
        chapter_para = review_doc.add_paragraph(style='List Bullet')
        chapter_para.add_run(f"Paragraphs: {stats['paragraphs']}")
        chapter_para.add_run(f"\nWords: {stats['words']}")
        chapter_para.add_run(f"\nAverage paragraph length: {stats['words'] // max(1, stats['paragraphs'])} words")
    
    # Add simple recommendations
    review_doc.add_heading("Basic Recommendations", level=2)
    
    recommendations = [
        "Review chapter lengths for consistency",
        "Check for any unusually short or long chapters",
        "Ensure chapter titles are descriptive and consistent",
        "Review the document for structural coherence",
        "Consider adding a table of contents for better navigation"
    ]
    
    for rec in recommendations:
        review_doc.add_paragraph(rec, style='List Bullet')
    
    # Save the review document
    try:
        os.makedirs(app.output_dir.get(), exist_ok=True)
        review_filepath = os.path.join(app.output_dir.get(), "Document_Review_Report.docx")
        review_doc.save(review_filepath)
        
        app.log(f"Local review document saved to: {review_filepath}")
        app.update_progress(100, "Local review completed")
        
        messagebox.showinfo("Fallback Review", 
            f"A basic document review has been generated as a fallback.\n\n"
            f"The review has been saved to: {review_filepath}"
        )
        
        return True
        
    except Exception as e:
        app.log(f"Error saving local review document: {str(e)}")
        app.update_progress(0, "Error saving review document")
        raise
