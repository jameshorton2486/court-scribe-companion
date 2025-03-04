
import os
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from tkinter import messagebox

def save_all_chapters(app):
    if not app.chapters:
        messagebox.showerror("Error", "No chapters available to save.")
        return
    
    try:
        app.log("Saving all chapters...")
        app.update_progress(0, "Saving chapters...")
        
        # Create output directory if it doesn't exist
        os.makedirs(app.output_dir.get(), exist_ok=True)
        
        total_chapters = len(app.chapters)
        for i, chapter in enumerate(app.chapters):
            app.update_progress((i / total_chapters) * 100, f"Saving chapter {i+1} of {total_chapters}")
            
            # Create a new document for the chapter
            doc = Document()
            
            # Add a header with book title and author
            header = doc.sections[0].header
            header_para = header.paragraphs[0]
            header_para.text = f"{app.book_title.get()} - {app.author_name.get()}"
            
            # Add the chapter title
            title = doc.add_heading(chapter['title'], level=1)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add the chapter content
            for para in chapter['content'][1:]:  # Skip the title paragraph
                if para.style.name.startswith('Heading'):
                    level = int(para.style.name[-1])
                    doc.add_heading(para.text, level=level)
                else:
                    doc.add_paragraph(para.text)
            
            # Create a safe filename from the chapter title
            safe_title = re.sub(r'[^\w\s-]', '', chapter['title']).strip().replace(' ', '_')
            chapter_filename = f"Chapter_{i+1}_{safe_title}.docx"
            chapter_path = os.path.join(app.output_dir.get(), chapter_filename)
            
            # Save the document
            doc.save(chapter_path)
            app.log(f"Saved chapter: {chapter_filename}")
        
        app.update_progress(100, "All chapters saved")
        app.log("All chapters saved successfully")
        messagebox.showinfo("Success", f"All chapters saved to {app.output_dir.get()}")
        
    except Exception as e:
        app.log(f"Error saving chapters: {str(e)}")
        app.update_progress(0, "Error saving chapters")
        messagebox.showerror("Error", f"Failed to save chapters: {str(e)}")

def generate_complete_book(app):
    if not app.chapters:
        messagebox.showerror("Error", "No chapters available to generate book.")
        return
    
    try:
        app.log("Generating complete book...")
        app.update_progress(0, "Generating book...")
        
        # Create output directory if it doesn't exist
        os.makedirs(app.output_dir.get(), exist_ok=True)
        
        # Create a new document
        doc = Document()
        
        # Add title page
        title = doc.add_heading(app.book_title.get(), level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        author = doc.add_paragraph(f"By {app.author_name.get()}")
        author.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add page break
        doc.add_page_break()
        
        # Add table of contents if available
        if app.toc:
            app.update_progress(20, "Adding table of contents...")
            doc.add_heading("Table of Contents", level=1)
            
            for item in app.toc:
                toc_text = f"{'    ' * (item['level'] - 1)}{item['title']}"
                toc_para = doc.add_paragraph(toc_text)
                
                # Add tab and page number (this will be a placeholder)
                # In a real TOC, this would be linked to actual page numbers
                toc_para.add_run("\t")
                toc_para.add_run("p. #")
            
            doc.add_page_break()
        
        # Add each chapter
        app.update_progress(40, "Adding chapters...")
        total_chapters = len(app.chapters)
        
        for i, chapter in enumerate(app.chapters):
            app.update_progress(40 + (i / total_chapters) * 50, f"Adding chapter {i+1} of {total_chapters}")
            
            # Add chapter title
            title = doc.add_heading(chapter['title'], level=1)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add chapter content
            for para in chapter['content'][1:]:  # Skip the title paragraph
                if para.style.name.startswith('Heading'):
                    level = int(para.style.name[-1])
                    doc.add_heading(para.text, level=level)
                else:
                    doc.add_paragraph(para.text)
            
            # Add page break between chapters
            if i < total_chapters - 1:
                doc.add_page_break()
        
        # Set document properties
        doc.core_properties.title = app.book_title.get()
        doc.core_properties.author = app.author_name.get()
        
        # Save the document
        book_filename = f"{app.book_title.get().replace(' ', '_')}_Complete.docx"
        book_path = os.path.join(app.output_dir.get(), book_filename)
        
        app.update_progress(90, "Saving complete book...")
        doc.save(book_path)
        
        app.update_progress(100, "Book generated successfully")
        app.log(f"Complete book saved to: {book_path}")
        messagebox.showinfo("Success", f"Complete book generated and saved to: {book_path}")
        
    except Exception as e:
        app.log(f"Error generating book: {str(e)}")
        app.update_progress(0, "Error generating book")
        messagebox.showerror("Error", f"Failed to generate book: {str(e)}")
