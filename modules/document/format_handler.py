
from docx import Document
from bs4 import BeautifulSoup
import re
import os

def detect_file_format(file_path):
    """Detect the format of a file based on its extension."""
    _, ext = os.path.splitext(file_path)
    ext = ext.lower()
    
    if ext == '.docx':
        return 'docx'
    elif ext == '.txt':
        return 'txt'
    elif ext == '.md':
        return 'md'
    elif ext == '.html' or ext == '.htm':
        return 'html'
    else:
        return 'unknown'

def load_docx_document(file_path):
    """Load a DOCX document and return a Document object."""
    return Document(file_path)

def load_text_document(file_path, encoding='utf-8'):
    """Load a text document and return its content."""
    try:
        with open(file_path, 'r', encoding=encoding) as f:
            return f.read()
    except UnicodeDecodeError:
        # If UTF-8 fails, try with latin-1 (which can read any byte sequence)
        with open(file_path, 'r', encoding='latin-1') as f:
            return f.read()

def load_html_document(file_path, encoding='utf-8'):
    """Load an HTML document and return a BeautifulSoup object."""
    try:
        with open(file_path, 'r', encoding=encoding) as f:
            content = f.read()
        return BeautifulSoup(content, 'html.parser')
    except UnicodeDecodeError:
        # If UTF-8 fails, try with latin-1
        with open(file_path, 'r', encoding='latin-1') as f:
            content = f.read()
        return BeautifulSoup(content, 'html.parser')

def convert_html_to_document(html_soup):
    """Convert an HTML BeautifulSoup object to a Document object."""
    doc = Document()
    
    # Process headings
    for i in range(1, 7):
        for heading in html_soup.find_all(f'h{i}'):
            doc.add_heading(heading.get_text(), level=i)
            
    # Process paragraphs
    for para in html_soup.find_all('p'):
        doc.add_paragraph(para.get_text())
        
    # Process lists
    for ul in html_soup.find_all('ul'):
        for li in ul.find_all('li'):
            doc.add_paragraph(f"• {li.get_text()}", style='List Bullet')
            
    for ol in html_soup.find_all('ol'):
        for i, li in enumerate(ol.find_all('li')):
            doc.add_paragraph(f"{i+1}. {li.get_text()}", style='List Number')
    
    return doc

def convert_markdown_to_document(markdown_content):
    """Convert Markdown content to a Document object."""
    doc = Document()
    
    lines = markdown_content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        # Process headings
        heading_match = re.match(r'^(#{1,6})\s+(.+)$', line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            doc.add_heading(text, level=level)
        
        # Process paragraphs (collect consecutive lines)
        elif line and not line.startswith('```'):
            paragraph_text = [line]
            j = i + 1
            while j < len(lines) and lines[j].strip() and not lines[j].strip().startswith('#') and not lines[j].strip().startswith('```'):
                paragraph_text.append(lines[j].strip())
                j += 1
            
            # Join the paragraph lines and add to document
            full_paragraph = ' '.join(paragraph_text)
            
            # Check if it's a list item
            if full_paragraph.startswith('* ') or full_paragraph.startswith('- '):
                doc.add_paragraph(full_paragraph[2:], style='List Bullet')
            elif re.match(r'^\d+\.\s', full_paragraph):
                doc.add_paragraph(full_paragraph, style='List Number')
            else:
                doc.add_paragraph(full_paragraph)
            
            i = j - 1  # -1 because of the i+=1 at the end of the loop
        
        # Skip code blocks
        elif line.startswith('```'):
            j = i + 1
            while j < len(lines) and not lines[j].strip() == '```':
                j += 1
            i = j
        
        i += 1
    
    return doc

def extract_chapters_from_headings(doc, min_heading_level=1, max_heading_level=2):
    """Extract chapters from a document based on heading levels."""
    chapters = []
    current_chapter = None
    chapter_content = []
    
    for para in doc.paragraphs:
        # Check if paragraph is a heading of the specified levels
        if para.style.name.startswith('Heading') or para.style.name.startswith('Title'):
            # Get heading level
            if para.style.name.startswith('Heading'):
                level = int(para.style.name[-1])
            else:  # Title
                level = 1
                
            # If the heading is within our chapter level range
            if min_heading_level <= level <= max_heading_level:
                # If we already have a chapter in progress, save it
                if current_chapter:
                    chapters.append({
                        'title': current_chapter,
                        'content': chapter_content
                    })
                
                # Start a new chapter
                current_chapter = para.text
                chapter_content = [para]
            else:
                # Add the heading to the current chapter
                if current_chapter:
                    chapter_content.append(para)
        else:
            # Add to current chapter if one exists, otherwise ignore
            if current_chapter:
                chapter_content.append(para)
    
    # Add the last chapter if it exists
    if current_chapter:
        chapters.append({
            'title': current_chapter,
            'content': chapter_content
        })
    
    return chapters

def detect_chapter_patterns(doc):
    """Heuristically detect chapter patterns in a document."""
    potential_chapter_starts = []
    
    # Look for common chapter patterns in the first 30 paragraphs
    chapter_patterns = [
        r'^chapter\s+\d+',  # "Chapter 1"
        r'^section\s+\d+',  # "Section 1"
        r'^part\s+\d+',     # "Part 1"
        r'^\d+\.\s',        # "1. Chapter Title"
        r'^chapter\s+[\w\s]+:',  # "Chapter One: The Beginning"
    ]
    
    sample_paragraphs = doc.paragraphs[:min(30, len(doc.paragraphs))]
    
    for i, para in enumerate(sample_paragraphs):
        for pattern in chapter_patterns:
            if re.match(pattern, para.text.lower()):
                potential_chapter_starts.append(i)
                break
    
    return potential_chapter_starts
