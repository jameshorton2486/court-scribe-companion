
import gc
from docx import Document

def convert_large_markdown_to_document(content, app, chunk_size=1000):
    """Optimized conversion for large markdown content."""
    app.log.info("Using optimized conversion for large markdown content")
    
    # Create a new document
    doc = Document()
    
    # Split content into lines for chunk processing
    lines = content.split('\n')
    total_lines = len(lines)
    
    # Process in chunks
    for i in range(0, total_lines, chunk_size):
        end_idx = min(i + chunk_size, total_lines)
        chunk = lines[i:end_idx]
        
        # Join chunk back into text
        chunk_text = '\n'.join(chunk)
        
        # Convert chunk
        from modules.document.format_handler import convert_markdown_to_document
        temp_doc = convert_markdown_to_document(chunk_text)
        
        # Add paragraphs from temp_doc to main doc
        for para in temp_doc.paragraphs:
            p = doc.add_paragraph()
            p.text = para.text
            p.style = para.style
        
        # Update progress
        progress = 20 + (i / total_lines) * 30
        app.update_progress(progress, f"Converting markdown content ({i}/{total_lines} lines)...")
        
        # Periodically yield to UI to prevent freezing
        app.update()
        
        # Force garbage collection between chunks
        if i % (chunk_size * 3) == 0:
            gc.collect()
    
    return doc
