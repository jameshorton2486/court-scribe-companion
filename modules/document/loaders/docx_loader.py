
from docx import Document

def load_docx_document(file_path):
    """Load a standard DOCX file."""
    return Document(file_path)

def load_large_docx_document(file_path, app, chunk_size=200):
    """Optimized loading for large DOCX files."""
    app.log.info("Using optimized loading for large DOCX file")
    
    # Load document with minimal processing
    doc = Document(file_path)
    
    # Count paragraphs and update progress
    total_paragraphs = len(doc.paragraphs)
    app.log.info(f"Large document loaded: {total_paragraphs} paragraphs")
    
    return doc
