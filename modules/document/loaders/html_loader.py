
import gc
from bs4 import BeautifulSoup
from docx import Document

def load_html_document(file_path):
    """Load an HTML file as BeautifulSoup object."""
    with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
        html_content = f.read()
    return BeautifulSoup(html_content, 'html.parser')

def convert_large_html_to_document(soup, app, chunk_size=100):
    """Optimized conversion for large HTML content."""
    app.log.info("Using optimized conversion for large HTML content")
    
    # Create a new document
    doc = Document()
    
    # Find all content blocks
    content_blocks = soup.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'ul', 'ol'])
    total_blocks = len(content_blocks)
    
    # Process in chunks
    for i in range(0, total_blocks, chunk_size):
        end_idx = min(i + chunk_size, total_blocks)
        chunk = content_blocks[i:end_idx]
        
        # Process each element in the chunk
        for element in chunk:
            if element.name.startswith('h') and len(element.name) == 2:
                # Heading element
                level = int(element.name[1])
                doc.add_heading(element.get_text(), level=level)
            elif element.name == 'p':
                # Paragraph element
                doc.add_paragraph(element.get_text())
            elif element.name in ['ul', 'ol']:
                # List element
                for li in element.find_all('li'):
                    doc.add_paragraph(li.get_text(), style='List Bullet' if element.name == 'ul' else 'List Number')
        
        # Update progress
        progress = 20 + (i / total_blocks) * 30
        app.update_progress(progress, f"Converting HTML content ({i}/{total_blocks} blocks)...")
        
        # Periodically yield to UI to prevent freezing
        app.update()
        
        # Force garbage collection between chunks
        if i % (chunk_size * 3) == 0:
            gc.collect()
    
    return doc
