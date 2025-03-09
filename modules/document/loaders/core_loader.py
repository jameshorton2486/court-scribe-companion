
import os
import gc
import time
import psutil
from docx import Document
from tkinter import messagebox
from modules.document.format_handler import detect_file_format

def load_document(app):
    """Main function to load document from file."""
    if not app.input_file.get():
        messagebox.showerror("Error", "Please select an input file")
        return
    
    try:
        file_path = app.input_file.get()
        app.log.info(f"Loading document: {file_path}")
        app.update_progress(10, "Loading document...")
        
        # Get file size for optimization decisions
        file_size = os.path.getsize(file_path)
        file_size_mb = file_size / (1024 * 1024)
        is_large_file = file_size_mb > 10  # Consider files > 10MB as large
        
        # Get available system memory
        available_memory_mb = psutil.virtual_memory().available / (1024 * 1024)
        memory_critical = available_memory_mb < 500  # Critical if < 500MB available
        
        if is_large_file:
            app.log.info(f"Large document detected: {file_size_mb:.2f} MB. Using optimized loading.")
            app.log.info(f"Available memory: {available_memory_mb:.2f} MB")
            
            # Clear memory before loading large file
            gc.collect()
        
        start_time = time.time()
        
        # Detect file format
        file_format = detect_file_format(file_path)
        app.log.info(f"Detected file format: {file_format}")
        
        # Set chunk sizes based on available memory
        if memory_critical:
            text_chunk_size = 500  # Smaller chunks for low memory
            docx_chunk_size = 50
        else:
            text_chunk_size = 1000  # Larger chunks for normal memory
            docx_chunk_size = 200
        
        # Import format-specific loaders
        from modules.document.loaders.docx_loader import load_docx_document, load_large_docx_document
        from modules.document.loaders.text_loader import load_text_document, load_large_text_document
        from modules.document.loaders.markdown_loader import convert_large_markdown_to_document
        from modules.document.loaders.html_loader import load_html_document, convert_large_html_to_document
        from modules.document.format_handler import convert_markdown_to_document, convert_html_to_document
        
        # Load document based on format
        if file_format == 'docx':
            # For large DOCX files, use optimized loading
            if is_large_file:
                doc = load_large_docx_document(file_path, app, docx_chunk_size)
            else:
                # Load the DOCX document normally
                doc = load_docx_document(file_path)
            app.docx_content = doc
            
        elif file_format == 'txt':
            # Load text file with optimizations for large files
            if is_large_file:
                content = load_large_text_document(file_path, app, text_chunk_size)
            else:
                content = load_text_document(file_path)
            
            # Create a Document object from text
            doc = Document()
            
            # For large text files, process in chunks
            if is_large_file:
                lines = content.split('\n')
                chunk_size = text_chunk_size
                total_lines = len(lines)
                
                for i in range(0, total_lines, chunk_size):
                    end_idx = min(i + chunk_size, total_lines)
                    chunk = lines[i:end_idx]
                    
                    for line in chunk:
                        if line.strip():  # Skip empty lines
                            doc.add_paragraph(line)
                    
                    # Update progress
                    progress = 10 + (i / total_lines) * 30
                    app.update_progress(progress, f"Processing text content ({i}/{total_lines} lines)...")
                    
                    # Force garbage collection between chunks
                    if i % (chunk_size * 3) == 0:
                        gc.collect()
            else:
                # Split the text content by lines and add as paragraphs
                lines = content.split('\n')
                for line in lines:
                    if line.strip():  # Skip empty lines
                        doc.add_paragraph(line)
            
            app.docx_content = doc
            
        elif file_format == 'md':
            # Load markdown file
            content = load_text_document(file_path)
            
            # Convert markdown to Document with optimizations for large files
            if is_large_file:
                doc = convert_large_markdown_to_document(content, app, text_chunk_size)
            else:
                doc = convert_markdown_to_document(content)
            app.docx_content = doc
            
        elif file_format == 'html':
            # Load HTML file
            soup = load_html_document(file_path)
            
            # Convert HTML to Document with optimizations for large files
            if is_large_file:
                doc = convert_large_html_to_document(soup, app, text_chunk_size)
            else:
                doc = convert_html_to_document(soup)
            app.docx_content = doc
            
        else:
            messagebox.showerror("Error", f"Unsupported file format: {file_format}")
            app.update_progress(0, "Error loading document")
            return
        
        # Calculate and log loading time
        loading_time = time.time() - start_time
        app.log.info(f"Document loaded in {loading_time:.2f} seconds")
        app.update_progress(50, f"Document loaded successfully in {loading_time:.2f}s")
        
        # Try to extract title from content or filename
        extract_document_title(app, doc, file_path)
        
        # For large files, force garbage collection after loading
        if is_large_file:
            gc.collect()
        
        app.update_progress(100, "Document loaded successfully")
        app.log.info(f"Document loaded successfully with {len(doc.paragraphs)} paragraphs")
        messagebox.showinfo("Success", "Document loaded successfully")
        
    except MemoryError:
        app.log.error("Memory error while loading document. The file may be too large.")
        messagebox.showerror("Memory Error", "Not enough memory to load this document. Try closing other applications or use a smaller file.")
        app.update_progress(0, "Memory error loading document")
    except Exception as e:
        app.log.error(f"Error loading document: {str(e)}")
        messagebox.showerror("Error", f"Failed to load document: {str(e)}")
        app.update_progress(0, "Error loading document")

def extract_document_title(app, doc, file_path):
    """Extract document title from content or filename."""
    if not app.book_title.get():
        # First check if there's a Title style paragraph
        title_found = False
        for para in doc.paragraphs[:10]:  # Only check first 10 paragraphs for efficiency
            if para.style.name.startswith('Title') or para.style.name.startswith('Heading 1'):
                app.book_title.set(para.text)
                title_found = True
                break
        
        # If still no title, use filename
        if not title_found:
            filename = os.path.basename(file_path)
            title, _ = os.path.splitext(filename)
            app.book_title.set(title)
