
import threading
import tkinter as tk
from tkinter import messagebox
from docx import Document
import os
from modules.document.format_handler import (
    detect_file_format, load_docx_document, load_text_document,
    load_html_document, convert_html_to_document, convert_markdown_to_document,
    extract_chapters_from_headings, detect_chapter_patterns
)
from modules.document.text_processor import (
    fix_text_encoding, clean_html_content, extract_text_from_markdown, normalize_whitespace
)

def load_document(app):
    if not app.input_file.get():
        messagebox.showerror("Error", "Please select an input file")
        return
    
    try:
        file_path = app.input_file.get()
        app.log(f"Loading document: {file_path}")
        app.update_progress(10, "Loading document...")
        
        # Detect file format
        file_format = detect_file_format(file_path)
        app.log(f"Detected file format: {file_format}")
        
        # Load document based on format
        if file_format == 'docx':
            # Load the DOCX document
            doc = load_docx_document(file_path)
            app.docx_content = doc
            
        elif file_format == 'txt':
            # Load text file
            content = load_text_document(file_path)
            
            # Create a Document object from text
            doc = Document()
            
            # Split the text content by lines and add as paragraphs
            lines = content.split('\n')
            for line in lines:
                if line.strip():  # Skip empty lines
                    para = doc.add_paragraph(line)
            
            app.docx_content = doc
            
        elif file_format == 'md':
            # Load markdown file
            content = load_text_document(file_path)
            
            # Convert markdown to Document
            doc = convert_markdown_to_document(content)
            app.docx_content = doc
            
        elif file_format == 'html':
            # Load HTML file
            soup = load_html_document(file_path)
            
            # Convert HTML to Document
            doc = convert_html_to_document(soup)
            app.docx_content = doc
            
        else:
            messagebox.showerror("Error", f"Unsupported file format: {file_format}")
            app.update_progress(0, "Error loading document")
            return
        
        app.update_progress(50, "Document loaded successfully")
        
        # Try to extract title from content or filename
        if not app.book_title.get():
            # First check if there's a Title style paragraph
            for para in doc.paragraphs[:10]:
                if para.style.name.startswith('Title') or para.style.name.startswith('Heading 1'):
                    app.book_title.set(para.text)
                    break
            
            # If still no title, use filename
            if not app.book_title.get():
                filename = os.path.basename(file_path)
                title, _ = os.path.splitext(filename)
                app.book_title.set(title)
        
        app.update_progress(100, "Document loaded successfully")
        app.log(f"Document loaded successfully with {len(doc.paragraphs)} paragraphs")
        messagebox.showinfo("Success", "Document loaded successfully")
        
    except Exception as e:
        app.log(f"Error loading document: {str(e)}")
        messagebox.showerror("Error", f"Failed to load document: {str(e)}")
        app.update_progress(0, "Error loading document")

def process_document(app):
    if not app.docx_content:
        messagebox.showerror("Error", "Please load a document first")
        return
    
    # If batch processing is enabled, process all files
    if app.batch_process.get() and len(app.input_files) > 1:
        threading.Thread(target=_batch_process_documents, args=(app,), daemon=True).start()
    else:
        # Process single document
        threading.Thread(target=_process_document_thread, args=(app,), daemon=True).start()

def _batch_process_documents(app):
    # ... keep existing code (batch processing logic)

def _process_document_thread(app, show_message=True):
    from modules.document.content_enhancer import enhance_book_content
    
    try:
        app.log("Starting document processing...")
        app.update_progress(0, "Processing document...")
        
        # Check if output directory exists, create if not
        os.makedirs(app.output_dir.get(), exist_ok=True)
        
        app.log(f"Processing document: {app.input_file.get()}")
        app.log(f"Output directory: {app.output_dir.get()}")
        app.log(f"Book title: {app.book_title.get()}")
        app.log(f"Author: {app.author_name.get()}")
        
        # Fix encoding if selected
        if app.fix_encoding.get():
            app.update_progress(20, "Fixing text encoding...")
            fix_text_encoding(app)
        
        # Extract chapters using multiple detection methods
        app.update_progress(40, "Extracting chapters...")
        app.chapters = []
        
        # First try to extract based on heading styles
        heading_chapters = extract_chapters_from_headings(
            app.docx_content,
            min_heading_level=1,
            max_heading_level=2
        )
        
        if heading_chapters:
            app.log(f"Found {len(heading_chapters)} chapters based on headings")
            app.chapters = heading_chapters
        else:
            app.log("No chapters found by heading analysis, trying pattern detection...")
            
            # Try to detect chapter patterns
            chapter_indices = detect_chapter_patterns(app.docx_content)
            
            if chapter_indices:
                app.log(f"Found {len(chapter_indices)} potential chapter patterns")
                
                # Convert chapter indices to chapters
                for i, idx in enumerate(chapter_indices):
                    # Determine end index (start of next chapter or end of document)
                    end_idx = chapter_indices[i+1] if i+1 < len(chapter_indices) else len(app.docx_content.paragraphs)
                    
                    chapter_title = app.docx_content.paragraphs[idx].text
                    chapter_content = app.docx_content.paragraphs[idx:end_idx]
                    
                    app.chapters.append({
                        'title': chapter_title,
                        'content': chapter_content
                    })
            else:
                # If no patterns found, try to split by blank lines
                app.log("No chapter patterns found, splitting by logical sections...")
                
                current_chapter = None
                chapter_content = []
                section_break_count = 0
                
                for i, para in enumerate(app.docx_content.paragraphs):
                    # Check if this paragraph is empty (potential section break)
                    if not para.text.strip():
                        section_break_count += 1
                    else:
                        # If we have consecutive blank lines and text after, potential new section
                        if section_break_count >= 2 and i > 0:
                            # If we have content already, save as chapter
                            if current_chapter and chapter_content:
                                app.chapters.append({
                                    'title': current_chapter,
                                    'content': chapter_content
                                })
                                
                                # Start new chapter with this paragraph as title
                                current_chapter = para.text
                                chapter_content = [para]
                            else:
                                # First chapter
                                current_chapter = para.text
                                chapter_content = [para]
                        else:
                            # Continue current chapter
                            if current_chapter:
                                chapter_content.append(para)
                            else:
                                # First paragraph becomes chapter title
                                current_chapter = para.text
                                chapter_content = [para]
                        
                        # Reset section break counter
                        section_break_count = 0
                
                # Add the last chapter if it exists
                if current_chapter and chapter_content:
                    app.chapters.append({
                        'title': current_chapter,
                        'content': chapter_content
                    })
        
        # If no chapters were found by any method, create a single chapter
        if not app.chapters:
            app.log("No chapters found with any method, creating a single chapter...")
            app.chapters.append({
                'title': app.book_title.get() or "Untitled Chapter",
                'content': app.docx_content.paragraphs
            })
        
        app.log(f"Extracted {len(app.chapters)} chapters")
        
        # Generate table of contents if selected
        if app.generate_toc.get():
            app.update_progress(60, "Generating table of contents...")
            app.toc = []
            
            for i, chapter in enumerate(app.chapters):
                app.toc.append({
                    'title': chapter['title'],
                    'level': 1,
                    'index': i
                })
                
                # Look for subheadings within the chapter
                for para in chapter['content']:
                    if para.style.name.startswith('Heading 2'):
                        app.toc.append({
                            'title': para.text,
                            'level': 2,
                            'index': i
                        })
                    elif para.style.name.startswith('Heading 3'):
                        app.toc.append({
                            'title': para.text,
                            'level': 3,
                            'index': i
                        })
            
            app.log(f"Generated table of contents with {len(app.toc)} entries")
        
        # Enhance content if selected
        if app.enhance_content.get():
            app.update_progress(80, "Enhancing book content...")
            enhance_book_content(app)
        
        # Update the chapter listbox
        app.chapter_listbox.delete(0, tk.END)
        for i, chapter in enumerate(app.chapters):
            app.chapter_listbox.insert(tk.END, f"Chapter {i+1}: {chapter['title']}")
        
        app.update_progress(100, "Document processing completed")
        app.log("Document processing completed successfully")
        
        if show_message:
            messagebox.showinfo("Success", "Document processing completed successfully")
        
    except Exception as e:
        app.log(f"Error processing document: {str(e)}")
        app.update_progress(0, "Error processing document")
        if show_message:
            messagebox.showerror("Error", f"Failed to process document: {str(e)}")
